"""
Convert Veltrix Sports documentation to Word document WITH EMBEDDED IMAGES
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Complete Project Documentation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.add_run('Version: ').bold = True
    info.add_run('1.0\n')
    info.add_run('Date: ').bold = True
    info.add_run('August 29, 2026\n')
    info.add_run('Status: ').bold = True
    info.add_run('Documentation Complete')
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('TABLE OF CONTENTS', 1)
    toc_items = [
        '1. Project Overview',
        '2. System Architecture',
        '3. Database Design',
        '4. API Documentation',
        '5. Feature Specifications',
        '6. Authentication Flow',
        '7. Training Plan Module',
        '8. Events Module',
        '9. Ticketing Module',
        '10. Payment Flow',
        '11. Device Integration',
        '12. Screen Navigation',
        '13. Deployment Guide',
        '14. Testing Strategy',
        '15. Security Checklist',
        '16. Performance Optimization',
        '17. Team Workflow',
        '18. Risk Assessment',
        '19. Cost Breakdown',
        '20. Diagrams Index'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. PROJECT OVERVIEW', 1)
    
    doc.add_heading('1.1 Project Summary', 2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ('Project Name', 'Veltrix Sports'),
        ('Type', 'Coaching, Events & Ticketing Platform'),
        ('Platforms', 'Android, iOS, Web'),
        ('Timeline', '2-12 weeks (MVP to Full Product)'),
        ('Team Size', '5-10 developers'),
        ('Budget', '₹43,750 - ₹56,28,000')
    ]
    for i, (key, value) in enumerate(data):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph('')
    
    doc.add_heading('1.2 Three Core Verticals', 2)
    doc.add_paragraph('Training Plan (B2B & B2C): Running, Cycling, Triathlon plans with AI generation and device sync')
    doc.add_paragraph('Events: Event listing, registration, ticketing, live tracking, and results')
    doc.add_paragraph('Ticketing Marketplace: Self-serve organizer, public marketplace, QR check-in, commission system')
    
    doc.add_heading('1.3 Technical Stack', 2)
    stack_table = doc.add_table(rows=8, cols=2)
    stack_table.style = 'Table Grid'
    stack_data = [
        ('Layer', 'Technology'),
        ('Frontend', 'Flutter 3.41.9, Dart 3.11.5'),
        ('State Management', 'BLoC'),
        ('Navigation', 'GoRouter'),
        ('Backend', 'Node.js/Express'),
        ('Database', 'PostgreSQL (AWS RDS)'),
        ('Cache', 'Redis (ElastiCache)'),
        ('Payments', 'Razorpay')
    ]
    for i, (key, value) in enumerate(stack_data):
        stack_table.cell(i, 0).text = key
        stack_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. System Architecture
    doc.add_heading('2. SYSTEM ARCHITECTURE', 1)
    doc.add_paragraph('The system follows a layered architecture with the following components:')
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '01_system_architecture.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.1 Architecture Layers', 2)
    layers = [
        'Client Layer: Android App, iOS App, Web App (Flutter)',
        'Gateway Layer: AWS API Gateway with rate limiting',
        'Service Layer: Auth Service (Cognito), Main API (Node.js), Payment Service (Razorpay)',
        'Data Layer: PostgreSQL (RDS), Redis Cache (ElastiCache), S3 Storage',
        'External Services: Firebase (Notifications), Garmin/Strava (Device Sync), AWS SES (Email)'
    ]
    for layer in layers:
        doc.add_paragraph(layer, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. Database Design
    doc.add_heading('3. DATABASE DESIGN', 1)
    doc.add_paragraph('The database consists of 24 tables with the following key entities:')
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '02_database_er.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.1 Core Tables', 2)
    tables_list = [
        'users - User accounts and profiles',
        'training_plans - Workout plans',
        'plan_sessions - Individual workout sessions',
        'session_logs - Completed session records',
        'events - Sports events',
        'event_categories - Event types/distances',
        'event_registrations - Event signups',
        'tickets - Purchased tickets',
        'devices - Connected fitness devices',
        'payments - Transaction records',
        'notifications - User alerts',
        'reviews - Event ratings'
    ]
    for table_item in tables_list:
        doc.add_paragraph(table_item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. API Documentation
    doc.add_heading('4. API DOCUMENTATION', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '09_api_endpoints.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('4.1 Base Configuration', 2)
    config_items = [
        'Base URL: https://api.veltrixsports.com/v1',
        'Protocol: HTTPS',
        'Format: JSON',
        'Authentication: Bearer Token (JWT)',
        'Timeout: 30 seconds',
        'Rate Limit: 100 requests/minute'
    ]
    for item in config_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.2 Authentication Endpoints', 2)
    auth_endpoints = [
        'POST /auth/register - Register new user',
        'POST /auth/login - User login',
        'POST /auth/otp/send - Send OTP to phone',
        'POST /auth/otp/verify - Verify OTP',
        'POST /auth/refresh - Refresh JWT token',
        'POST /auth/forgot-password - Reset password'
    ]
    for endpoint in auth_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_heading('4.3 Training Plan Endpoints', 2)
    plan_endpoints = [
        'GET /plans - Get user plans',
        'POST /plans/generate - Create AI plan',
        'GET /plans/{planId} - Get plan details',
        'PUT /plans/{planId} - Update plan',
        'DELETE /plans/{planId} - Delete plan',
        'POST /sessions/{sessionId}/complete - Log session'
    ]
    for endpoint in plan_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_heading('4.4 Event Endpoints', 2)
    event_endpoints = [
        'GET /events - List events',
        'GET /events/{eventId} - Get event details',
        'POST /events/{eventId}/register - Register for event',
        'GET /events/registrations - Get user registrations'
    ]
    for endpoint in event_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_heading('4.5 Payment Endpoints', 2)
    payment_endpoints = [
        'POST /payments/create - Create payment order',
        'POST /payments/verify - Verify payment',
        'GET /payments - Get payment history'
    ]
    for endpoint in payment_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Feature Specifications
    doc.add_heading('5. FEATURE SPECIFICATIONS', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '08_feature_modules.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('5.1 Feature Summary', 2)
    feature_table = doc.add_table(rows=7, cols=4)
    feature_table.style = 'Table Grid'
    feature_data = [
        ('Module', 'Features', 'P0', 'P1'),
        ('Authentication', '4', '3', '1'),
        ('Training Plan', '6', '4', '2'),
        ('Events', '4', '3', '1'),
        ('Ticketing', '4', '0', '3'),
        ('Profile', '4', '0', '4'),
        ('Total', '26', '10', '11')
    ]
    for i, row_data in enumerate(feature_data):
        for j, cell_data in enumerate(row_data):
            feature_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 6. Authentication Flow
    doc.add_heading('6. AUTHENTICATION FLOW', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '03_authentication_flow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('The authentication flow includes:')
    
    auth_steps = [
        'User opens app → Splash screen',
        'Decision: Has account?',
        'No → Register Screen → Enter details → Send OTP → Verify → Account created',
        'Yes → Login Screen → Enter credentials → Validate → API call → Store token → Dashboard',
        'Security: JWT tokens, password hashing, account lockout, OTP expiry'
    ]
    for step in auth_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 7. Training Plan Module
    doc.add_heading('7. TRAINING PLAN MODULE', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '04_training_plan_flow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('7.1 Plan Creation Flow', 2)
    plan_steps = [
        'Select Sport → Running, Cycling, Triathlon',
        'Select Goal → 5K, 10K, Half Marathon, Marathon',
        'AI Generate or Manual?',
        'Review Plan → Preview weekly schedule',
        'Save Plan → Add to user\'s plans'
    ]
    for step in plan_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('7.2 Session Completion', 2)
    session_steps = [
        'Select Session → View details',
        'Start Session → Begin tracking',
        'Complete Session → Log actual metrics',
        'Log Results → Save to database'
    ]
    for step in session_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 8. Events Module
    doc.add_heading('8. EVENTS MODULE', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '05_event_registration_flow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('8.1 Event Discovery', 2)
    event_steps = [
        'Browse Events → View event list',
        'Filter/Search → By sport, city, date',
        'Select Event → View details',
        'Register → Proceed to registration'
    ]
    for step in event_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('8.2 Registration Process', 2)
    reg_steps = [
        'Select Category → Choose distance/type',
        'Fill Form → Personal information',
        'Emergency Contact → Required for safety',
        'Medical Info → Blood group, conditions',
        'Proceed to Payment → Complete registration'
    ]
    for step in reg_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 9. Ticketing Module
    doc.add_heading('9. TICKETING MODULE', 1)
    doc.add_paragraph('The ticketing module includes:')
    
    ticketing_features = [
        'Marketplace: Browse available tickets',
        'Filter by city/sport',
        'Sort by price/date',
        'Purchase Flow: Select quantity → Apply promo → Payment → Confirmation',
        'QR Check-in: Camera scanning, ticket validation'
    ]
    for feature in ticketing_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Payment Flow
    doc.add_heading('10. PAYMENT FLOW', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '06_payment_flow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('10.1 Payment Process', 2)
    payment_steps = [
        'User proceeds to payment',
        'Create Razorpay order',
        'Open Razorpay checkout',
        'Select payment method',
        'Process payment',
        'Verify signature',
        'Capture payment',
        'Send confirmation',
        'Generate ticket'
    ]
    for step in payment_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('10.2 Security', 2)
    security_items = [
        'PCI DSS compliant (via Razorpay)',
        'Signature verification',
        'Idempotency keys',
        'No card details stored'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 11. Device Integration
    doc.add_heading('11. DEVICE INTEGRATION', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '07_device_sync_flow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('11.1 Supported Devices', 2)
    device_table = doc.add_table(rows=5, cols=3)
    device_table.style = 'Table Grid'
    device_data = [
        ('Device', 'Provider', 'Data Synced'),
        ('Garmin', 'Garmin Connect', 'Activities, Heart Rate'),
        ('Apple Watch', 'Apple Health', 'Workouts, Health Data'),
        ('Strava', 'Strava API', 'Activities, Routes'),
        ('Wear OS', 'Google Fit', 'Activities, Health Data')
    ]
    for i, row_data in enumerate(device_data):
        for j, cell_data in enumerate(row_data):
            device_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 12. Screen Navigation
    doc.add_heading('12. SCREEN NAVIGATION', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '10_screen_navigation.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('The app has 15+ screens organized as:')
    
    screen_list = [
        'Auth: Splash, Login, Register, OTP',
        'Main: Dashboard, Training Plans, Events, Tickets, Profile',
        'Detail: Plan Detail, Session Detail, Event Detail, Ticket Detail, Settings',
        'Bottom Navigation: 5 main tabs'
    ]
    for screen in screen_list:
        doc.add_paragraph(screen, style='List Bullet')
    
    doc.add_page_break()
    
    # 13. Deployment Guide
    doc.add_heading('13. DEPLOYMENT GUIDE', 1)
    
    # EMBED IMAGE
    img_path = os.path.join(os.path.dirname(__file__), 'diagrams', '12_deployment_architecture.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('13.1 AWS Services', 2)
    aws_table = doc.add_table(rows=10, cols=3)
    aws_table.style = 'Table Grid'
    aws_data = [
        ('Service', 'Purpose', 'Config'),
        ('CloudFront', 'CDN', 'Global'),
        ('ALB', 'Load Balancing', 'Multi-AZ'),
        ('EC2', 'Compute', 't3.small × 2'),
        ('RDS', 'Database', 'PostgreSQL'),
        ('ElastiCache', 'Caching', 'Redis'),
        ('S3', 'Storage', 'Encrypted'),
        ('Cognito', 'Auth', 'User pools'),
        ('Razorpay', 'Payments', 'PCI Compliant'),
        ('Firebase', 'Notifications', 'FCM')
    ]
    for i, row_data in enumerate(aws_data):
        for j, cell_data in enumerate(row_data):
            aws_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 14. Testing Strategy
    doc.add_heading('14. TESTING STRATEGY', 1)
    
    doc.add_heading('14.1 Test Pyramid', 2)
    test_items = [
        'Unit Tests (70%): BLoC, Repository, Use Case, Model tests',
        'Integration Tests (20%): API, Database, Firebase integration',
        'E2E Tests (10%): Login, Registration, Training, Event flows'
    ]
    for item in test_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('14.2 Coverage Target', 2)
    doc.add_paragraph('80% minimum code coverage required')
    
    doc.add_page_break()
    
    # 15. Security Checklist
    doc.add_heading('15. SECURITY CHECKLIST', 1)
    
    doc.add_heading('15.1 Authentication', 2)
    auth_security = [
        'OAuth 2.0 implemented',
        'JWT tokens used',
        'Token expiry configured',
        'Secure password hashing',
        'Account lockout enabled'
    ]
    for item in auth_security:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('15.2 Data Protection', 2)
    data_security = [
        'HTTPS enforced',
        'Database encryption',
        'S3 encryption enabled',
        'Sensitive data encrypted',
        'No secrets in code'
    ]
    for item in data_security:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('15.3 API Security', 2)
    api_security = [
        'Input validation',
        'SQL injection prevented',
        'XSS attacks prevented',
        'Rate limiting enabled',
        'CORS configured'
    ]
    for item in api_security:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_page_break()
    
    # 16. Performance Optimization
    doc.add_heading('16. PERFORMANCE OPTIMIZATION', 1)
    
    doc.add_heading('16.1 Targets', 2)
    perf_table = doc.add_table(rows=5, cols=2)
    perf_table.style = 'Table Grid'
    perf_data = [
        ('Metric', 'Target'),
        ('App Load Time', '< 2 seconds'),
        ('API Response Time', '< 500ms'),
        ('Frame Rate', '60 FPS'),
        ('Memory Usage', '< 200MB')
    ]
    for i, (key, value) in enumerate(perf_data):
        perf_table.cell(i, 0).text = key
        perf_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 17. Team Workflow
    doc.add_heading('17. TEAM WORKFLOW', 1)
    
    doc.add_heading('17.1 Agile Process', 2)
    agile_items = [
        'Sprint Duration: 2 weeks',
        'Daily Standup: 9:30 AM',
        'Sprint Planning: Monday, 10:00 AM',
        'Sprint Review: Friday, 3:00 PM',
        'Retrospective: Friday, 4:00 PM'
    ]
    for item in agile_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('17.2 Git Strategy', 2)
    git_items = [
        'main → production',
        'develop → integration',
        'feature/* → new features',
        'bugfix/* → bug fixes',
        'hotfix/* → urgent fixes'
    ]
    for item in git_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 18. Risk Assessment
    doc.add_heading('18. RISK ASSESSMENT', 1)
    
    risk_table = doc.add_table(rows=6, cols=5)
    risk_table.style = 'Table Grid'
    risk_data = [
        ('Risk', 'Probability', 'Impact', 'Score', 'Mitigation'),
        ('Scope Creep', 'High', 'High', '8/10', 'Clear MVP scope'),
        ('Timeline Delays', 'High', 'High', '8/10', 'Buffer time'),
        ('API Integration', 'High', 'Medium', '6/10', 'Abstraction layer'),
        ('Performance', 'Medium', 'High', '6/10', 'Optimization'),
        ('Security', 'Low', 'Critical', '4/10', 'Security audits')
    ]
    for i, row_data in enumerate(risk_data):
        for j, cell_data in enumerate(row_data):
            risk_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 19. Cost Breakdown
    doc.add_heading('19. COST BREAKDOWN', 1)
    
    doc.add_heading('19.1 MVP (2 Weeks)', 2)
    mvp_table = doc.add_table(rows=4, cols=2)
    mvp_table.style = 'Table Grid'
    mvp_data = [
        ('Item', 'Cost'),
        ('Development Tools', '₹0'),
        ('AWS (Free Tier)', '₹0'),
        ('Team (5 people)', '₹43,750-87,500')
    ]
    for i, (key, value) in enumerate(mvp_data):
        mvp_table.cell(i, 0).text = key
        mvp_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('19.2 Full Product (3 Months)', 2)
    full_table = doc.add_table(rows=4, cols=2)
    full_table.style = 'Table Grid'
    full_data = [
        ('Item', 'Cost'),
        ('Development Tools', '₹8,700'),
        ('AWS Services', '₹24,600'),
        ('Team (9 people)', '₹4,95,000-9,90,000')
    ]
    for i, (key, value) in enumerate(full_data):
        full_table.cell(i, 0).text = key
        full_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 20. Diagrams Index
    doc.add_heading('20. ALL DIAGRAMS', 1)
    doc.add_paragraph('Below are all 12 diagrams embedded in this document:')
    
    diagrams = [
        ('01_system_architecture.png', 'System Architecture'),
        ('02_database_er.png', 'Database ER Diagram'),
        ('03_authentication_flow.png', 'Authentication Flow'),
        ('04_training_plan_flow.png', 'Training Plan Flow'),
        ('05_event_registration_flow.png', 'Event Registration Flow'),
        ('06_payment_flow.png', 'Payment Flow'),
        ('07_device_sync_flow.png', 'Device Sync Flow'),
        ('08_feature_modules.png', 'Feature Modules'),
        ('09_api_endpoints.png', 'API Endpoints'),
        ('10_screen_navigation.png', 'Screen Navigation'),
        ('11_data_flow.png', 'Data Flow'),
        ('12_deployment_architecture.png', 'Deployment Architecture')
    ]
    
    for filename, title in diagrams:
        doc.add_heading(title, 2)
        img_path = os.path.join(os.path.dirname(__file__), 'diagrams', filename)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_COMPLETE.docx')
    doc.save(output_path)
    print(f"Document saved with embedded images: {output_path}")
    return output_path

if __name__ == '__main__':
    create_document()
