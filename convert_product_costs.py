"""
Convert product costs to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_product_costs():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Product Launch & Scaling Costs', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.add_run('Date: ').bold = True
    info.add_run('August 29, 2026\n')
    info.add_run('Purpose: ').bold = True
    info.add_run('Production & Scaling Budget')
    
    doc.add_page_break()
    
    # 1. ONE-TIME COSTS
    doc.add_heading('1. ONE-TIME COSTS (Launch)', 1)
    
    doc.add_heading('1.1 App Store Fees', 2)
    appstore_table = doc.add_table(rows=4, cols=3)
    appstore_table.style = 'Table Grid'
    appstore_data = [
        ('Store', 'Cost', 'Notes'),
        ('Google Play Developer', '₹18,000', 'One-time, lifetime'),
        ('Apple Developer', '₹7,500', 'Annual renewal'),
        ('Total', '₹25,500', '')
    ]
    for i, row_data in enumerate(appstore_data):
        for j, cell_data in enumerate(row_data):
            appstore_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('1.2 Domain & SSL', 2)
    domain_table = doc.add_table(rows=3, cols=3)
    domain_table.style = 'Table Grid'
    domain_data = [
        ('Item', 'Cost', 'Duration'),
        ('Domain (.com)', '₹800', '1 year'),
        ('SSL Certificate', '₹0', 'Free (Let\'s Encrypt)')
    ]
    for i, row_data in enumerate(domain_data):
        for j, cell_data in enumerate(row_data):
            domain_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('1.3 Business Registration', 2)
    biz_table = doc.add_table(rows=4, cols=2)
    biz_table.style = 'Table Grid'
    biz_data = [
        ('Item', 'Cost'),
        ('GST Registration', '₹0 (Free)'),
        ('MSME Registration', '₹0 (Free)'),
        ('Trademark (Optional)', '₹5,000')
    ]
    for i, (key, value) in enumerate(biz_data):
        biz_table.cell(i, 0).text = key
        biz_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. MONTHLY COSTS
    doc.add_heading('2. MONTHLY COSTS (Production)', 1)
    
    doc.add_heading('2.1 Infrastructure (AWS)', 2)
    aws_table = doc.add_table(rows=6, cols=3)
    aws_table.style = 'Table Grid'
    aws_data = [
        ('Service', 'Specification', 'Monthly Cost'),
        ('EC2 (Compute)', 't3.small', '₹1,400'),
        ('RDS (Database)', 'db.t3.small, 20GB', '₹2,800'),
        ('S3 (Storage)', '10GB', '₹200'),
        ('CloudFront (CDN)', '100GB transfer', '₹800'),
        ('Total', '', '₹5,300')
    ]
    for i, row_data in enumerate(aws_data):
        for j, cell_data in enumerate(row_data):
            aws_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('2.2 Scaling Costs', 2)
    scale_table = doc.add_table(rows=6, cols=3)
    scale_table.style = 'Table Grid'
    scale_data = [
        ('Users', 'Infrastructure', 'Monthly Cost'),
        ('0-1,000', 't3.small', '₹5,300'),
        ('1,000-5,000', 't3.medium', '₹12,000'),
        ('5,000-10,000', 't3.large + Read Replica', '₹25,000'),
        ('10,000-50,000', 'm5.large + Multi-AZ', '₹60,000'),
        ('50,000+', 'Cluster setup', '₹1,50,000+')
    ]
    for i, row_data in enumerate(scale_data):
        for j, cell_data in enumerate(row_data):
            scale_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('2.3 Third-Party Services', 2)
    third_table = doc.add_table(rows=5, cols=3)
    third_table.style = 'Table Grid'
    third_data = [
        ('Service', 'Cost/Month', 'Notes'),
        ('Razorpay', '2% of transactions', 'Pay per use'),
        ('Firebase', '₹0', 'Free tier'),
        ('Twilio (SMS)', '₹1,500', '~500 OTPs/month'),
        ('Total', '~₹1,500 + 2%', '')
    ]
    for i, row_data in enumerate(third_data):
        for j, cell_data in enumerate(row_data):
            third_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 3. TOTAL MONTHLY COST
    doc.add_heading('3. TOTAL MONTHLY COST', 1)
    
    doc.add_heading('Phase 1: Launch (0-1,000 users)', 2)
    phase1_table = doc.add_table(rows=4, cols=2)
    phase1_table.style = 'Table Grid'
    phase1_data = [
        ('Category', 'Cost'),
        ('Infrastructure', '₹5,300'),
        ('Third-party', '₹1,500'),
        ('Total', '₹6,900/month')
    ]
    for i, (key, value) in enumerate(phase1_data):
        phase1_table.cell(i, 0).text = key
        phase1_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Phase 2: Growth (1,000-5,000 users)', 2)
    phase2_table = doc.add_table(rows=5, cols=2)
    phase2_table.style = 'Table Grid'
    phase2_data = [
        ('Category', 'Cost'),
        ('Infrastructure', '₹12,000'),
        ('Third-party', '₹5,000'),
        ('Support + Marketing', '₹35,000'),
        ('Total', '₹52,000/month')
    ]
    for i, (key, value) in enumerate(phase2_data):
        phase2_table.cell(i, 0).text = key
        phase2_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Phase 3: Scale (5,000-10,000 users)', 2)
    phase3_table = doc.add_table(rows=5, cols=2)
    phase3_table.style = 'Table Grid'
    phase3_data = [
        ('Category', 'Cost'),
        ('Infrastructure', '₹25,000'),
        ('Third-party', '₹15,000'),
        ('Support + Marketing', '₹80,000'),
        ('Total', '₹1,20,000/month')
    ]
    for i, (key, value) in enumerate(phase3_data):
        phase3_table.cell(i, 0).text = key
        phase3_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 4. REVENUE MODEL
    doc.add_heading('4. REVENUE MODEL', 1)
    
    doc.add_heading('4.1 Revenue Streams', 2)
    revenue_table = doc.add_table(rows=5, cols=3)
    revenue_table.style = 'Table Grid'
    revenue_data = [
        ('Stream', 'Pricing', 'Revenue/User'),
        ('Event Registration', '5-10% commission', '₹50-200'),
        ('Ticket Sales', '10-15% commission', '₹100-500'),
        ('Premium Plans', '₹499-999/month', '₹499-999'),
        ('Ads (Future)', 'CPM', '₹10-50')
    ]
    for i, row_data in enumerate(revenue_data):
        for j, cell_data in enumerate(row_data):
            revenue_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('4.2 Break-Even Analysis', 2)
    breakeven_table = doc.add_table(rows=5, cols=3)
    breakeven_table.style = 'Table Grid'
    breakeven_data = [
        ('Phase', 'Revenue Needed', 'Users Needed'),
        ('Phase 1', '₹6,900', '50 registrations'),
        ('Phase 2', '₹52,000', '400 registrations'),
        ('Phase 3', '₹1,20,000', '800 registrations'),
        ('Phase 4', '₹2,65,000', '1,500 registrations')
    ]
    for i, row_data in enumerate(breakeven_data):
        for j, cell_data in enumerate(row_data):
            breakeven_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 5. SCALING PLAN
    doc.add_heading('5. SCALING PLAN', 1)
    
    phases = [
        ('Phase 1: MVP (Month 1-2)', '0-1,000 users', '₹6,900/month', 'Product-market fit'),
        ('Phase 2: Growth (Month 3-6)', '1,000-5,000 users', '₹52,000/month', 'User acquisition'),
        ('Phase 3: Scale (Month 7-12)', '5,000-10,000 users', '₹1,20,000/month', 'Retention & revenue'),
        ('Phase 4: Enterprise (Year 2+)', '10,000+ users', '₹2,65,000+/month', 'Market leadership')
    ]
    
    for phase_name, users, cost, focus in phases:
        doc.add_heading(phase_name, 2)
        doc.add_paragraph(f'• Users: {users}')
        doc.add_paragraph(f'• Cost: {cost}')
        doc.add_paragraph(f'• Focus: {focus}')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    # 6. BUDGET SUMMARY
    doc.add_heading('6. BUDGET SUMMARY', 1)
    
    doc.add_heading('One-Time Costs', 2)
    onetime_table = doc.add_table(rows=4, cols=2)
    onetime_table.style = 'Table Grid'
    onetime_data = [
        ('Item', 'Cost'),
        ('App Store Fees', '₹25,500'),
        ('Domain', '₹800'),
        ('Business Registration', '₹5,000')
    ]
    for i, (key, value) in enumerate(onetime_data):
        onetime_table.cell(i, 0).text = key
        onetime_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Monthly Costs (Start)', 2)
    monthly_table = doc.add_table(rows=3, cols=2)
    monthly_table.style = 'Table Grid'
    monthly_data = [
        ('Item', 'Cost'),
        ('Infrastructure', '₹5,300'),
        ('Third-party', '₹1,500')
    ]
    for i, (key, value) in enumerate(monthly_data):
        monthly_table.cell(i, 0).text = key
        monthly_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Year 1 Total', 2)
    year1_table = doc.add_table(rows=5, cols=2)
    year1_table.style = 'Table Grid'
    year1_data = [
        ('Phase', 'Cost'),
        ('One-Time', '₹31,300'),
        ('Phase 1 (2 months)', '₹13,800'),
        ('Phase 2 (4 months)', '₹2,08,000'),
        ('Phase 3 (6 months)', '₹7,20,000')
    ]
    for i, (key, value) in enumerate(year1_data):
        year1_table.cell(i, 0).text = key
        year1_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 7. INVESTMENT REQUIRED
    doc.add_heading('7. INVESTMENT REQUIRED', 1)
    
    doc.add_heading('Minimum Viable Product', 2)
    mvp_table = doc.add_table(rows=3, cols=2)
    mvp_table.style = 'Table Grid'
    mvp_data = [
        ('Item', 'Amount'),
        ('One-time costs', '₹31,300'),
        ('3 months runway', '₹20,700')
    ]
    for i, (key, value) in enumerate(mvp_data):
        mvp_table.cell(i, 0).text = key
        mvp_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Total MVP Investment: ₹52,000').bold = True
    
    doc.add_paragraph('')
    
    doc.add_heading('12 Months Operations', 2)
    year_table = doc.add_table(rows=4, cols=2)
    year_table.style = 'Table Grid'
    year_data = [
        ('Item', 'Amount'),
        ('One-time costs', '₹31,300'),
        ('12 months operations', '₹82,800'),
        ('Buffer (20%)', '₹22,800')
    ]
    for i, (key, value) in enumerate(year_data):
        year_table.cell(i, 0).text = key
        year_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Total 12-Month Investment: ₹1,36,900').bold = True
    
    doc.add_page_break()
    
    # 8. COST OPTIMIZATION
    doc.add_heading('8. COST OPTIMIZATION', 1)
    
    doc.add_heading('Free Tier Usage', 2)
    free_items = [
        'AWS Free Tier (12 months): Save ₹63,600',
        'Firebase Free Tier: Save ₹0',
        'Free SSL: Save ₹5,000/year'
    ]
    for item in free_items:
        doc.add_paragraph('✓ ' + item)
    
    doc.add_heading('Total Savings', 2)
    savings_items = [
        'Year 1 savings: ~₹70,000',
        'Actual Year 1 cost: ~₹9,00,000'
    ]
    for item in savings_items:
        doc.add_paragraph('• ' + item)
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_PRODUCT_COSTS.docx')
    doc.save(output_path)
    print(f"Product costs document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_product_costs()
