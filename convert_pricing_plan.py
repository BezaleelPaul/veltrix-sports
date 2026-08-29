"""
Convert pricing and application plan to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_pricing_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Pricing & Application Plan', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.add_run('Date: ').bold = True
    info.add_run('August 29, 2026\n')
    info.add_run('Status: ').bold = True
    info.add_run('Ready for Review')
    
    doc.add_page_break()
    
    # 1. PRICING BREAKDOWN
    doc.add_heading('1. PRICING BREAKDOWN', 1)
    
    # Option A
    doc.add_heading('1.1 Option A: MVP (2 Weeks) - RECOMMENDED', 2)
    
    dev_table = doc.add_table(rows=5, cols=5)
    dev_table.style = 'Table Grid'
    dev_data = [
        ('Role', 'Count', 'Rate/Day', 'Days', 'Total'),
        ('Flutter Developer', '2', '₹2,000', '10', '₹40,000'),
        ('Backend Developer', '1', '₹2,500', '10', '₹25,000'),
        ('Subtotal', '', '', '', '₹65,000'),
        ('TOTAL', '', '', '', '₹65,000')
    ]
    for i, row_data in enumerate(dev_data):
        for j, cell_data in enumerate(row_data):
            dev_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('UI/UX is in-house, so no additional cost. AWS Free Tier for 12 months.')
    
    doc.add_paragraph('')
    
    # Option B
    doc.add_heading('1.2 Option B: Full Product (1 Month)', 2)
    
    full_table = doc.add_table(rows=6, cols=5)
    full_table.style = 'Table Grid'
    full_data = [
        ('Role', 'Count', 'Rate/Day', 'Days', 'Total'),
        ('Flutter Developer', '2', '₹2,000', '22', '₹88,000'),
        ('Backend Developer', '1', '₹2,500', '22', '₹55,000'),
        ('QA Engineer', '1', '₹1,500', '15', '₹22,500'),
        ('Subtotal', '', '', '', '₹1,65,500'),
        ('TOTAL', '', '', '', '₹1,70,500')
    ]
    for i, row_data in enumerate(full_data):
        for j, cell_data in enumerate(row_data):
            full_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Option C
    doc.add_heading('1.3 Option C: Complete Product (3 Months)', 2)
    
    ent_table = doc.add_table(rows=8, cols=5)
    ent_table.style = 'Table Grid'
    ent_data = [
        ('Role', 'Count', 'Rate/Day', 'Days', 'Total'),
        ('Flutter Developer', '3', '₹2,000', '66', '₹3,96,000'),
        ('Backend Developer', '2', '₹2,500', '66', '₹3,30,000'),
        ('QA Engineer', '1', '₹1,500', '44', '₹66,000'),
        ('DevOps', '1', '₹2,000', '22', '₹44,000'),
        ('Subtotal', '', '', '', '₹8,36,000'),
        ('AWS (3 months)', '', '', '', '₹15,000'),
        ('TOTAL', '', '', '', '₹8,51,000')
    ]
    for i, row_data in enumerate(ent_data):
        for j, cell_data in enumerate(row_data):
            ent_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # Monthly Running Cost
    doc.add_heading('1.4 Monthly Running Cost', 2)
    
    monthly_table = doc.add_table(rows=6, cols=2)
    monthly_table.style = 'Table Grid'
    monthly_data = [
        ('Service', 'Cost/Month'),
        ('AWS EC2', '₹1,400'),
        ('AWS RDS', '₹2,800'),
        ('AWS S3', '₹200'),
        ('AWS CloudFront', '₹800'),
        ('Total', '₹5,300/month')
    ]
    for i, (key, value) in enumerate(monthly_data):
        monthly_table.cell(i, 0).text = key
        monthly_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Third-Party Costs
    doc.add_heading('1.5 Third-Party Service Costs', 2)
    
    third_table = doc.add_table(rows=6, cols=2)
    third_table.style = 'Table Grid'
    third_data = [
        ('Service', 'Pricing'),
        ('Razorpay', '2% per transaction'),
        ('Firebase', 'Free (1K notifications/day)'),
        ('Garmin API', 'Free'),
        ('Apple Developer', '₹7,500/year'),
        ('Google Play', '₹25,000 (one-time)')
    ]
    for i, (key, value) in enumerate(third_data):
        third_table.cell(i, 0).text = key
        third_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. UI/UX DELIVERABLES
    doc.add_heading('2. UI/UX DELIVERABLES', 1)
    
    doc.add_heading('2.1 Design System', 2)
    design_system = [
        'Color palette (Primary, Secondary, Accent)',
        'Typography (Font family, sizes, weights)',
        'Spacing system (Margins, padding)',
        'Border radius',
        'Shadows',
        'Icons style',
        'Button styles',
        'Input field styles',
        'Card styles'
    ]
    for item in design_system:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('2.2 Screens Required (15+)', 2)
    
    doc.add_heading('Auth Screens', 3)
    auth_screens = [
        'Splash Screen',
        'Login Screen',
        'Register Screen',
        'OTP Verification Screen',
        'Forgot Password Screen'
    ]
    for screen in auth_screens:
        doc.add_paragraph('☐ ' + screen)
    
    doc.add_heading('Main Screens', 3)
    main_screens = [
        'Dashboard Screen',
        'Training Plans List Screen',
        'Plan Detail Screen',
        'Session Detail Screen',
        'Events List Screen',
        'Event Detail Screen',
        'Registration Form Screen',
        'My Registrations Screen',
        'Profile Screen',
        'Settings Screen'
    ]
    for screen in main_screens:
        doc.add_paragraph('☐ ' + screen)
    
    doc.add_heading('Additional Screens', 3)
    additional_screens = [
        'Payment Checkout Screen',
        'Payment Success Screen',
        'Payment Failure Screen',
        'Error States (404, 500)',
        'Empty States',
        'Loading States'
    ]
    for screen in additional_screens:
        doc.add_paragraph('☐ ' + screen)
    
    doc.add_heading('2.3 Components Required', 2)
    components = [
        'App Bar',
        'Bottom Navigation',
        'Side Drawer',
        'Cards (Event, Plan, Session)',
        'Buttons (Primary, Secondary, Text)',
        'Input Fields',
        'Dropdowns',
        'Badges',
        'Avatars',
        'Progress Indicators',
        'Modals/Bottom Sheets',
        'Toast/Snackbar'
    ]
    for component in components:
        doc.add_paragraph('☐ ' + component)
    
    doc.add_heading('2.4 Prototypes Required', 2)
    prototypes = [
        'Login flow prototype',
        'Registration flow prototype',
        'Dashboard flow prototype',
        'Training plan flow prototype',
        'Event registration flow prototype',
        'Payment flow prototype'
    ]
    for prototype in prototypes:
        doc.add_paragraph('☐ ' + prototype)
    
    doc.add_page_break()
    
    # 3. FEATURE LIST
    doc.add_heading('3. FEATURE LIST (FROM PDF)', 1)
    
    # Authentication
    doc.add_heading('3.1 Authentication Module', 2)
    auth_table = doc.add_table(rows=9, cols=3)
    auth_table.style = 'Table Grid'
    auth_data = [
        ('Feature', 'Priority', 'MVP'),
        ('User Registration (Email)', 'P0', 'Yes'),
        ('User Registration (Phone)', 'P0', 'Yes'),
        ('User Login', 'P0', 'Yes'),
        ('OTP Verification', 'P0', 'Yes'),
        ('Forgot Password', 'P1', 'No'),
        ('Social Login (Google)', 'P1', 'No'),
        ('Social Login (Apple)', 'P1', 'No'),
        ('Remember Me', 'P2', 'No')
    ]
    for i, row_data in enumerate(auth_data):
        for j, cell_data in enumerate(row_data):
            auth_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Training Plans
    doc.add_heading('3.2 Training Plans Module', 2)
    plan_table = doc.add_table(rows=9, cols=3)
    plan_table.style = 'Table Grid'
    plan_data = [
        ('Feature', 'Priority', 'MVP'),
        ('View Training Plans', 'P0', 'Yes'),
        ('Plan Detail View', 'P0', 'Yes'),
        ('Session Detail View', 'P0', 'Yes'),
        ('Complete Session', 'P0', 'Yes'),
        ('Create Training Plan', 'P1', 'No'),
        ('AI Plan Generation', 'P2', 'No'),
        ('Progress Dashboard', 'P1', 'No'),
        ('Training Calendar', 'P2', 'No')
    ]
    for i, row_data in enumerate(plan_data):
        for j, cell_data in enumerate(row_data):
            plan_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Events
    doc.add_heading('3.3 Events Module', 2)
    event_table = doc.add_table(rows=9, cols=3)
    event_table.style = 'Table Grid'
    event_data = [
        ('Feature', 'Priority', 'MVP'),
        ('Browse Events', 'P0', 'Yes'),
        ('Event Detail View', 'P0', 'Yes'),
        ('Filter Events', 'P0', 'Yes'),
        ('Search Events', 'P1', 'No'),
        ('Event Registration', 'P0', 'Yes'),
        ('My Registrations', 'P1', 'No'),
        ('Event Reviews', 'P2', 'No'),
        ('Event Sharing', 'P2', 'No')
    ]
    for i, row_data in enumerate(event_data):
        for j, cell_data in enumerate(row_data):
            event_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Ticketing
    doc.add_heading('3.4 Ticketing Module', 2)
    ticket_table = doc.add_table(rows=8, cols=3)
    ticket_table.style = 'Table Grid'
    ticket_data = [
        ('Feature', 'Priority', 'MVP'),
        ('Browse Marketplace', 'P1', 'No'),
        ('Buy Tickets', 'P1', 'No'),
        ('View My Tickets', 'P1', 'No'),
        ('QR Code Display', 'P1', 'No'),
        ('QR Check-in', 'P2', 'No'),
        ('Promo Codes', 'P2', 'No'),
        ('Refund Request', 'P2', 'No')
    ]
    for i, row_data in enumerate(ticket_data):
        for j, cell_data in enumerate(row_data):
            ticket_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Payments
    doc.add_heading('3.5 Payments Module', 2)
    payment_table = doc.add_table(rows=7, cols=3)
    payment_table.style = 'Table Grid'
    payment_data = [
        ('Feature', 'Priority', 'MVP'),
        ('Razorpay Integration', 'P0', 'Yes'),
        ('UPI Payment', 'P0', 'Yes'),
        ('Card Payment', 'P0', 'Yes'),
        ('Net Banking', 'P1', 'No'),
        ('Payment History', 'P1', 'No'),
        ('Refund Processing', 'P2', 'No')
    ]
    for i, row_data in enumerate(payment_data):
        for j, cell_data in enumerate(row_data):
            payment_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 4. MVP SCOPE
    doc.add_heading('4. MVP SCOPE (2 WEEKS)', 1)
    
    doc.add_heading('4.1 Must Build (P0)', 2)
    mvp_items = [
        'User Registration',
        'User Login',
        'OTP Verification',
        'Dashboard',
        'Training Plans List',
        'Plan Detail',
        'Events List',
        'Event Detail',
        'Event Registration',
        'Basic Payment'
    ]
    for item in mvp_items:
        doc.add_paragraph('✓ ' + item)
    
    doc.add_heading('4.2 Skip for MVP', 2)
    skip_items = [
        'AI Plan Generation',
        'Device Sync',
        'Ticketing Marketplace',
        'Admin Panel',
        'Analytics'
    ]
    for item in skip_items:
        doc.add_paragraph('✗ ' + item)
    
    doc.add_page_break()
    
    # 5. DEVELOPMENT TIMELINE
    doc.add_heading('5. DEVELOPMENT TIMELINE', 1)
    
    doc.add_heading('Week 1: Foundation', 2)
    week1_table = doc.add_table(rows=6, cols=3)
    week1_table.style = 'Table Grid'
    week1_data = [
        ('Day', 'Backend', 'Frontend'),
        ('Day 1', 'Setup Node.js, PostgreSQL, Auth APIs', 'Flutter project setup'),
        ('Day 2', 'OTP, JWT, User APIs', 'API client, Auth BLoC'),
        ('Day 3', 'Plan APIs', 'Auth screens'),
        ('Day 4', 'Event APIs', 'Dashboard screen'),
        ('Day 5', 'Payment APIs', 'Training plans screens')
    ]
    for i, row_data in enumerate(week1_data):
        for j, cell_data in enumerate(row_data):
            week1_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('Week 2: Features', 2)
    week2_table = doc.add_table(rows=6, cols=3)
    week2_table.style = 'Table Grid'
    week2_data = [
        ('Day', 'Backend', 'Frontend'),
        ('Day 6', 'Testing, Bug fixes', 'Events screens'),
        ('Day 7', 'Documentation', 'Registration form'),
        ('Day 8', 'Deployment setup', 'Payment integration'),
        ('Day 9', 'AWS setup', 'Testing, Bug fixes'),
        ('Day 10', 'Final testing', 'Build, Deploy')
    ]
    for i, row_data in enumerate(week2_data):
        for j, cell_data in enumerate(row_data):
            week2_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 6. TEAM RESPONSIBILITIES
    doc.add_heading('6. TEAM RESPONSIBILITIES', 1)
    
    doc.add_heading('UI/UX Team', 2)
    uiux_tasks = [
        'Design system creation',
        'All screen designs (15+ screens)',
        'Prototypes (6 flows)',
        'Asset export (icons, images)',
        'Design handoff to developers'
    ]
    for task in uiux_tasks:
        doc.add_paragraph('☐ ' + task)
    
    doc.add_heading('Backend Team', 2)
    backend_tasks = [
        'API development (10+ endpoints)',
        'Database setup (PostgreSQL)',
        'Authentication (JWT, OTP)',
        'Payment integration (Razorpay)',
        'Deployment (AWS)'
    ]
    for task in backend_tasks:
        doc.add_paragraph('☐ ' + task)
    
    doc.add_heading('Frontend Team', 2)
    frontend_tasks = [
        'Flutter app development',
        'Screen implementation (15+ screens)',
        'State management (BLoC)',
        'API integration',
        'Testing on Android & iOS'
    ]
    for task in frontend_tasks:
        doc.add_paragraph('☐ ' + task)
    
    doc.add_page_break()
    
    # 7. SUCCESS CRITERIA
    doc.add_heading('7. SUCCESS CRITERIA', 1)
    
    success_items = [
        'User can register with email/phone',
        'User can login',
        'User can see dashboard',
        'User can browse training plans',
        'User can view plan details',
        'User can browse events',
        'User can view event details',
        'User can register for events',
        'User can make payment',
        'App works on Android & iOS'
    ]
    for item in success_items:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_page_break()
    
    # 8. APPROVAL
    doc.add_heading('8. APPROVAL', 1)
    
    approval_table = doc.add_table(rows=5, cols=2)
    approval_table.style = 'Table Grid'
    approval_data = [
        ('Role', 'Signature'),
        ('Project Manager', '_______________'),
        ('Tech Lead', '_______________'),
        ('UI/UX Lead', '_______________'),
        ('Stakeholder', '_______________')
    ]
    for i, (key, value) in enumerate(approval_data):
        approval_table.cell(i, 0).text = key
        approval_table.cell(i, 1).text = value
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_PRICING_PLAN.docx')
    doc.save(output_path)
    print(f"Pricing document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_pricing_document()
