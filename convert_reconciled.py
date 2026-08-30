"""
Convert reconciled document to Word
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_reconciled():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Reconciled Project Proposal — Stakeholder Approval', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Document Info
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Document Version', '2.0'),
        ('Date', 'August 29, 2026'),
        ('Classification', 'Confidential — For Internal Review'),
        ('Prepared For', 'Project Stakeholder'),
        ('Prepared By', 'Veltrix Sports Development Team'),
        ('Status', 'Pending Approval')
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    
    doc.add_paragraph(
        'Veltrix Sports is a sports platform connecting athletes, coaches, event organizers, '
        'and spectators through three verticals: Training Plans, Events, and Ticketing.'
    )
    
    doc.add_paragraph(
        'This document reconciles the Product Feature List (4-page PDF) with the Project Proposal '
        '(22-page DOCX) and classifies all features into prioritized phases.'
    )
    
    doc.add_heading('Key Finding', 2)
    doc.add_paragraph(
        'The Product Feature List contains significantly more functionality than the original '
        '2-week MVP proposal. The full feature set requires 3-6 months of development, not 2 weeks.'
    )
    
    doc.add_heading('Recommended Approach', 2)
    approach_table = doc.add_table(rows=5, cols=4)
    approach_table.style = 'Table Grid'
    approach_data = [
        ('Phase', 'Timeline', 'Scope', 'Budget'),
        ('Phase 1: MVP', '2 weeks', 'Core features only', '₹91,800'),
        ('Phase 2: Launch', 'Weeks 3-8', 'Complete product', '₹155,000'),
        ('Phase 3: Scale', 'Month 3+', 'Enterprise features', '₹210,000'),
        ('Total', '6 months', 'Full platform', '₹456,800')
    ]
    for i, row_data in enumerate(approach_data):
        for j, cell_data in enumerate(row_data):
            approach_table.cell(i, j).text = cell_data
    
    doc.add_heading('Decision Required', 2)
    doc.add_paragraph(
        'Stakeholder must confirm: Phase 1 only (₹91,800) or Phase 1+2 (₹246,800)?'
    )
    
    doc.add_page_break()
    
    # 2. Scope Reconciliation
    doc.add_heading('2. Scope Reconciliation Summary', 1)
    
    doc.add_heading('2.1 Feature Count Comparison', 2)
    compare_table = doc.add_table(rows=4, cols=3)
    compare_table.style = 'Table Grid'
    compare_data = [
        ('Source', 'Features', 'Estimated Screens'),
        ('Product Feature List PDF', '80+', '60-80'),
        ('Project Proposal DOCX', '28', '28'),
        ('Difference', '52+ features', '32-52 screens')
    ]
    for i, row_data in enumerate(compare_data):
        for j, cell_data in enumerate(row_data):
            compare_table.cell(i, j).text = cell_data
    
    doc.add_heading('2.2 Major Gaps Identified', 2)
    
    doc.add_heading('Training Plan Vertical', 3)
    training_gaps = [
        'AI plan generation — Not in proposal, P2 priority',
        'Wearable integration — Not in proposal, P1 priority',
        'Terrain integration — Not in proposal, P1 priority',
        'Weather integration — Not in proposal, P1 priority',
        'Zoom/Google Meet — Not in proposal, P1 priority',
        'In-app messaging — Not in proposal, P1 priority',
        'Subscription billing — Not in proposal, P1 priority',
        'Referral program — Not in proposal, P1 priority'
    ]
    for gap in training_gaps:
        doc.add_paragraph(gap, style='List Bullet')
    
    doc.add_heading('Events Vertical', 3)
    events_gaps = [
        'Organizer management — Not in proposal, P1 priority',
        'Bib allocation — Not in proposal, P1 priority',
        'Volunteer management — Not in proposal, P1 priority',
        'Route maps — Not in proposal, P1 priority',
        'Live tracking — Not in proposal, P1 priority',
        'Results/e-certificates — Not in proposal, P1 priority',
        'Sponsor management — Not in proposal, P1 priority',
        'Merchandise sales — Not in proposal, P1 priority'
    ]
    for gap in events_gaps:
        doc.add_paragraph(gap, style='List Bullet')
    
    doc.add_heading('Ticketing Vertical', 3)
    ticketing_gaps = [
        'Self-serve organizer dashboard — Not in proposal, P1 priority',
        'Multi-event marketplace — Not in proposal, P1 priority',
        'Custom registration forms — Not in proposal, P1 priority',
        'Commission/payout system — Not in proposal, P1 priority'
    ]
    for gap in ticketing_gaps:
        doc.add_paragraph(gap, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. Feature Prioritization
    doc.add_heading('3. Feature Prioritization', 1)
    
    doc.add_heading('3.1 P0 — Must-Have MVP (2-Week Scope)', 2)
    doc.add_paragraph('These features are essential for a working product that can be launched.')
    
    doc.add_heading('Training Plan (B2B & B2C)', 3)
    p0_training = [
        'UI/UX design',
        'User database',
        'Server, hosting, maintenance',
        'Domain',
        'Coach and athlete dashboards (basic)',
        'Landing pages (basic)',
        'Registration flow',
        'Welcome kit (digital)',
        'Payment gateway (Razorpay)',
        'Email/account setup',
        'Running plans',
        'Triathlon plans',
        'Biking plans',
        'Nutrition plan',
        'Hydration plan',
        'Strength & conditioning',
        'API integration',
        'User journey tracking',
        'Analytics (basic)'
    ]
    for item in p0_training:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Events', 3)
    p0_events = [
        'Event listing',
        'Event registration',
        'Tickets',
        'Payment gateway (Razorpay)',
        'API integration',
        'Analytics (basic)'
    ]
    for item in p0_events:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Ticketing', 3)
    p0_ticketing = [
        'Ticketing functionality',
        'QR code check-in',
        'Event reviews (basic)'
    ]
    for item in p0_ticketing:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Cross-Cutting', 3)
    p0_cross = [
        'Shared infrastructure',
        'API integrations',
        'Analytics (basic)'
    ]
    for item in p0_cross:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('3.2 P1 — Launch/Phase 2 (Weeks 3-8)', 2)
    doc.add_paragraph('These features are important for a complete product but can be added after initial launch.')
    
    doc.add_heading('Training Plan', 3)
    p1_training = [
        'Fitness/medical intake assessment',
        'Subscription/recurring billing',
        'Technique/education content library',
        'Device integration (wearables)',
        'Terrain integration',
        'Weather integration',
        'Zoom/Google Meet integration',
        'Performance/progress tracking',
        'In-app coach-athlete messaging',
        'Push notifications and reminders',
        'Referral program',
        'Customer support/helpdesk'
    ]
    for item in p1_training:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Events', 3)
    p1_events = [
        'Coaching plan add-on',
        'Team/relay registration',
        'Waitlist management',
        'Promo/discount codes',
        'Refund/cancellation rules',
        'Device integration',
        'Terrain integration',
        'Weather integration',
        'Organizer-side event creation tools',
        'Bib number allocation',
        'Volunteer management',
        'Route/course maps',
        'Live race tracking',
        'Results processing',
        'e-certificate generation',
        'Race photo marketplace',
        'Sponsor and branding management',
        'Merchandise sales',
        'Post-event feedback survey'
    ]
    for item in p1_events:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Ticketing', 3)
    p1_ticketing = [
        'Self-serve organizer dashboard',
        'Multi-event public marketplace',
        'Custom registration form builder',
        'Multiple ticket tiers and pricing',
        'Platform commission structure',
        'Automated payout/settlement',
        'SEO-optimized public event pages'
    ]
    for item in p1_ticketing:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('3.3 P2 — Future/Advanced (Month 3+)', 2)
    doc.add_paragraph('These features are advanced capabilities that require significant development.')
    
    doc.add_heading('Training Plan', 3)
    p2_training = [
        'AI automation of plan generation',
        'Advanced analytics dashboard',
        'White-label solutions'
    ]
    for item in p2_training:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Events', 3)
    p2_events = [
        'Advanced live tracking',
        'Advanced results processing',
        'Advanced sponsor management'
    ]
    for item in p2_events:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Ticketing', 3)
    p2_ticketing = [
        'Advanced marketplace features',
        'Dynamic pricing',
        'Advanced SEO'
    ]
    for item in p2_ticketing:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Cross-Cutting', 3)
    p2_cross = [
        'Admin/CMS panel',
        'Data privacy and compliance (GDPR, DPDP)',
        'Marketing tools (email/SMS campaigns)',
        'Business-side revenue dashboard',
        'Customer support/helpdesk'
    ]
    for item in p2_cross:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Revised Technical Architecture
    doc.add_heading('4. Revised Technical Architecture', 1)
    
    doc.add_heading('4.1 Architecture Layers', 2)
    arch_table = doc.add_table(rows=10, cols=3)
    arch_table.style = 'Table Grid'
    arch_data = [
        ('Layer', 'Technology', 'Purpose'),
        ('Frontend', 'Flutter 3.41.9', 'Cross-platform mobile'),
        ('Language', 'Dart 3.11.5', 'Application logic'),
        ('State Management', 'BLoC', 'Application state'),
        ('Navigation', 'GoRouter', 'Screen navigation'),
        ('HTTP Client', 'Dio', 'API communication'),
        ('Local Storage', 'Hive', 'Local persistence'),
        ('Backend', '[REQUIRES CONFIRMATION]', 'Backend services'),
        ('Database', '[REQUIRES CONFIRMATION]', 'Data persistence'),
        ('Payment', 'Razorpay', 'Payment processing')
    ]
    for i, row_data in enumerate(arch_data):
        for j, cell_data in enumerate(row_data):
            arch_table.cell(i, j).text = cell_data
    
    doc.add_heading('4.2 Backend Options', 2)
    doc.add_paragraph('[BACKEND ARCHITECTURE REQUIRES CONFIRMATION]')
    
    backend_table = doc.add_table(rows=4, cols=3)
    backend_table.style = 'Table Grid'
    backend_data = [
        ('Option', 'Services', 'Best For'),
        ('Firebase-Primary', 'Firebase Auth, Firestore, Storage', 'Quick MVP'),
        ('AWS-Primary', 'Cognito, RDS, S3', 'Scalability'),
        ('Hybrid', 'Firebase Auth, PostgreSQL, S3', 'Balance')
    ]
    for i, row_data in enumerate(backend_data):
        for j, cell_data in enumerate(row_data):
            backend_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 5. Revised Timeline
    doc.add_heading('5. Revised Timeline', 1)
    
    doc.add_heading('5.1 Phase 1: MVP (2 Weeks)', 2)
    phase1_table = doc.add_table(rows=11, cols=3)
    phase1_table.style = 'Table Grid'
    phase1_data = [
        ('Day', 'Focus', 'Deliverables'),
        ('1', 'Backend Setup', 'Database, Auth APIs, Scaffolding'),
        ('2', 'Flutter Setup', 'Structure, Theme, Navigation'),
        ('3', 'Authentication', 'Login, Signup, OTP'),
        ('4', 'Home Dashboard', 'Home, Quick actions'),
        ('5', 'Training Module', 'Plans, Details, Sessions'),
        ('6', 'Events Module', 'Events, Details, Registration'),
        ('7', 'Ticketing Module', 'Tickets, Seat selection'),
        ('8', 'Payments', 'Razorpay, Cart, Checkout'),
        ('9', 'Testing', 'Bug fixes, Polish'),
        ('10', 'Deployment', 'Build, Upload')
    ]
    for i, row_data in enumerate(phase1_data):
        for j, cell_data in enumerate(row_data):
            phase1_table.cell(i, j).text = cell_data
    
    doc.add_heading('5.2 Phase 2: Launch (Weeks 3-8)', 2)
    phase2_table = doc.add_table(rows=4, cols=3)
    phase2_table.style = 'Table Grid'
    phase2_data = [
        ('Week', 'Focus', 'Deliverables'),
        ('3-4', 'Wearables + Integrations', 'Device sync, Terrain, Weather'),
        ('5-6', 'Events Operations', 'Organizer tools, Bibs, Live tracking'),
        ('7-8', 'Marketplace', 'Self-serve dashboard, Commission system')
    ]
    for i, row_data in enumerate(phase2_data):
        for j, cell_data in enumerate(row_data):
            phase2_table.cell(i, j).text = cell_data
    
    doc.add_heading('5.3 Phase 3: Scale (Month 3+)', 2)
    phase3_table = doc.add_table(rows=4, cols=3)
    phase3_table.style = 'Table Grid'
    phase3_data = [
        ('Month', 'Focus', 'Deliverables'),
        ('3', 'AI + Advanced', 'AI plans, Advanced analytics'),
        ('4', 'Compliance + Marketing', 'GDPR, Marketing tools'),
        ('5-6', 'Enterprise', 'Admin CMS, Business dashboard')
    ]
    for i, row_data in enumerate(phase3_data):
        for j, cell_data in enumerate(row_data):
            phase3_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 6. Revised Cost Breakdown
    doc.add_heading('6. Revised Cost Breakdown', 1)
    
    doc.add_heading('6.1 Phase 1: MVP', 2)
    phase1_cost = doc.add_table(rows=6, cols=2)
    phase1_cost.style = 'Table Grid'
    phase1_cost_data = [
        ('Category', 'Amount (₹)'),
        ('Development', '65,000'),
        ('Infrastructure', '800'),
        ('Services', '500'),
        ('App Stores', '25,500'),
        ('Phase 1 Total', '91,800')
    ]
    for i, (key, value) in enumerate(phase1_cost_data):
        phase1_cost.cell(i, 0).text = key
        phase1_cost.cell(i, 1).text = value
    
    doc.add_heading('6.2 Phase 2: Launch', 2)
    phase2_cost = doc.add_table(rows=7, cols=2)
    phase2_cost.style = 'Table Grid'
    phase2_cost_data = [
        ('Category', 'Amount (₹)'),
        ('Flutter Development', '80,000'),
        ('Backend Development', '40,000'),
        ('UI/UX Design', '20,000'),
        ('Infrastructure', '5,000'),
        ('Services', '10,000'),
        ('Phase 2 Total', '155,000')
    ]
    for i, (key, value) in enumerate(phase2_cost_data):
        phase2_cost.cell(i, 0).text = key
        phase2_cost.cell(i, 1).text = value
    
    doc.add_heading('6.3 Phase 3: Scale', 2)
    phase3_cost = doc.add_table(rows=6, cols=2)
    phase3_cost.style = 'Table Grid'
    phase3_cost_data = [
        ('Category', 'Amount (₹)'),
        ('AI/ML Development', '100,000'),
        ('Backend Development', '60,000'),
        ('Infrastructure', '20,000'),
        ('Services', '30,000'),
        ('Phase 3 Total', '210,000')
    ]
    for i, (key, value) in enumerate(phase3_cost_data):
        phase3_cost.cell(i, 0).text = key
        phase3_cost.cell(i, 1).text = value
    
    doc.add_heading('6.4 Total Investment', 2)
    total_cost = doc.add_table(rows=5, cols=2)
    total_cost.style = 'Table Grid'
    total_cost_data = [
        ('Phase', 'Amount (₹)'),
        ('Phase 1: MVP', '91,800'),
        ('Phase 2: Launch', '155,000'),
        ('Phase 3: Scale', '210,000'),
        ('Grand Total', '456,800')
    ]
    for i, (key, value) in enumerate(total_cost_data):
        total_cost.cell(i, 0).text = key
        total_cost.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 7. Platform Scope
    doc.add_heading('7. Platform Scope', 1)
    platform_table = doc.add_table(rows=4, cols=4)
    platform_table.style = 'Table Grid'
    platform_data = [
        ('Platform', 'Phase 1', 'Phase 2', 'Phase 3'),
        ('Android', 'Primary', 'Primary', 'Primary'),
        ('iOS', '[REQUIRES CONFIRMATION]', 'Secondary', 'Secondary'),
        ('Web', 'No', '[REQUIRES CONFIRMATION]', '[REQUIRES CONFIRMATION]')
    ]
    for i, row_data in enumerate(platform_data):
        for j, cell_data in enumerate(row_data):
            platform_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 8. Team Requirements
    doc.add_heading('8. Team Requirements', 1)
    
    doc.add_heading('8.1 Phase 1 Team', 2)
    phase1_team = doc.add_table(rows=5, cols=2)
    phase1_team.style = 'Table Grid'
    phase1_team_data = [
        ('Role', 'Count'),
        ('Flutter Developer', '2'),
        ('Backend Developer', '1'),
        ('UI/UX Designer', '1'),
        ('QA Tester', '1')
    ]
    for i, (key, value) in enumerate(phase1_team_data):
        phase1_team.cell(i, 0).text = key
        phase1_team.cell(i, 1).text = value
    
    doc.add_heading('8.2 Phase 2 Team (Additional)', 2)
    phase2_team = doc.add_table(rows=4, cols=2)
    phase2_team.style = 'Table Grid'
    phase2_team_data = [
        ('Role', 'Count'),
        ('Flutter Developer', '+1'),
        ('Backend Developer', '+1'),
        ('DevOps Engineer', '1')
    ]
    for i, (key, value) in enumerate(phase2_team_data):
        phase2_team.cell(i, 0).text = key
        phase2_team.cell(i, 1).text = value
    
    doc.add_heading('8.3 Phase 3 Team (Additional)', 2)
    phase3_team = doc.add_table(rows=3, cols=2)
    phase3_team.style = 'Table Grid'
    phase3_team_data = [
        ('Role', 'Count'),
        ('AI/ML Engineer', '1'),
        ('Data Engineer', '1')
    ]
    for i, (key, value) in enumerate(phase3_team_data):
        phase3_team.cell(i, 0).text = key
        phase3_team.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 9. Acceptance Criteria
    doc.add_heading('9. Acceptance Criteria', 1)
    
    doc.add_heading('9.1 Phase 1 MVP Criteria', 2)
    phase1_ac = doc.add_table(rows=11, cols=2)
    phase1_ac.style = 'Table Grid'
    phase1_ac_data = [
        ('ID', 'Criterion'),
        ('AC-01', 'All 28 screens implemented'),
        ('AC-02', 'Authentication works (email, Google, Apple)'),
        ('AC-03', 'Training plans browsable and purchasable'),
        ('AC-04', 'Events browsable and registrable'),
        ('AC-05', 'Tickets purchasable with QR codes'),
        ('AC-06', 'Razorpay integration functional'),
        ('AC-07', 'Profile management works'),
        ('AC-08', 'All critical defects resolved'),
        ('AC-09', 'Release build passes QA'),
        ('AC-10', 'App store submission complete')
    ]
    for i, (key, value) in enumerate(phase1_ac_data):
        phase1_ac.cell(i, 0).text = key
        phase1_ac.cell(i, 1).text = value
    
    doc.add_heading('9.2 Phase 2 Launch Criteria', 2)
    phase2_ac = doc.add_table(rows=6, cols=2)
    phase2_ac.style = 'Table Grid'
    phase2_ac_data = [
        ('ID', 'Criterion'),
        ('AC-11', 'Wearable integration functional'),
        ('AC-12', 'Organizer tools available'),
        ('AC-13', 'Live tracking working'),
        ('AC-14', 'Marketplace operational'),
        ('AC-15', 'Commission system functional')
    ]
    for i, (key, value) in enumerate(phase2_ac_data):
        phase2_ac.cell(i, 0).text = key
        phase2_ac.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 10. Risks & Mitigation
    doc.add_heading('10. Risks & Mitigation', 1)
    risk_table = doc.add_table(rows=6, cols=4)
    risk_table.style = 'Table Grid'
    risk_data = [
        ('Risk', 'Probability', 'Impact', 'Mitigation'),
        ('Scope creep', 'High', 'High', 'Strict phase adherence'),
        ('Backend delays', 'Medium', 'High', 'Early API provisioning'),
        ('Payment issues', 'Medium', 'High', 'Early sandbox testing'),
        ('Timeline overrun', 'Medium', 'High', 'Daily standups'),
        ('Feature list expansion', 'High', 'Medium', 'Change request process')
    ]
    for i, row_data in enumerate(risk_data):
        for j, cell_data in enumerate(row_data):
            risk_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 11. Immediate Next Steps
    doc.add_heading('11. Immediate Next Steps', 1)
    next_table = doc.add_table(rows=7, cols=3)
    next_table.style = 'Table Grid'
    next_data = [
        ('#', 'Action', 'Owner'),
        ('1', 'Review this reconciled document', 'Stakeholder'),
        ('2', 'Confirm MVP scope (P0 only)', 'Stakeholder'),
        ('3', 'Confirm budget (₹91,800 or ₹246,800)', 'Stakeholder'),
        ('4', 'Confirm platform scope (Android/iOS)', 'Stakeholder'),
        ('5', 'Confirm backend architecture', 'Stakeholder + Tech'),
        ('6', 'Begin Phase 1 development', 'Team')
    ]
    for i, row_data in enumerate(next_data):
        for j, cell_data in enumerate(row_data):
            next_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 12. Conclusion
    doc.add_heading('12. Conclusion', 1)
    conclusion_table = doc.add_table(rows=7, cols=2)
    conclusion_table.style = 'Table Grid'
    conclusion_data = [
        ('Item', 'Status'),
        ('Scope Reconciliation', 'Complete'),
        ('Feature Prioritization', 'Complete'),
        ('Timeline', 'Defined (3 phases)'),
        ('Budget', 'Defined (3 phases)'),
        ('Backend Architecture', 'REQUIRES CONFIRMATION'),
        ('Platform Scope', 'REQUIRES CONFIRMATION')
    ]
    for i, (key, value) in enumerate(conclusion_data):
        conclusion_table.cell(i, 0).text = key
        conclusion_table.cell(i, 1).text = value
    
    doc.add_paragraph(
        'Decision Required: Approve Phase 1 only (₹91,800) or Phase 1+2 (₹246,800)?'
    )
    
    doc.add_page_break()
    
    # Appendix
    doc.add_heading('Appendix: Change Summary', 1)
    
    doc.add_heading('Critical Changes Made', 2)
    critical_changes = [
        'Scope reconciliation: Aligned Feature List PDF with Project Proposal',
        'Feature prioritization: Classified all features as P0/P1/P2',
        'Timeline revision: Expanded from 2 weeks to 3 phases (6 months)',
        'Budget revision: Expanded from ₹91,800 to ₹456,800 (3 phases)',
        'Platform scope: Marked iOS and Web as requiring confirmation',
        'Backend architecture: Marked as requiring confirmation'
    ]
    for change in critical_changes:
        doc.add_paragraph(change, style='List Bullet')
    
    doc.add_heading('Items Requiring Stakeholder Confirmation', 2)
    confirmation_items = [
        'MVP Scope: P0 only or P0+P1?',
        'Budget: ₹91,800 (Phase 1) or ₹246,800 (Phase 1+2)?',
        'Platform: Android only or Android+iOS?',
        'Backend: Firebase, AWS, or Hybrid?',
        'Ticketing: Basic or Marketplace?'
    ]
    for item in confirmation_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Risks Identified', 2)
    risks = [
        'Scope creep — High probability, high impact',
        'Feature list expansion — High probability, medium impact',
        'Backend delays — Medium probability, high impact',
        'Timeline overrun — Medium probability, high impact'
    ]
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_RECONCILED.docx')
    doc.save(output_path)
    print(f"Reconciled document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_reconciled()
