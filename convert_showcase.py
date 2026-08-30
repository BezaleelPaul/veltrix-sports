"""
Convert showcase document to Word
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_showcase():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Project Showcase Document', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_paragraph('')
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Item', 'Details'),
        ('Project', 'Veltrix Sports'),
        ('Type', 'Sports Platform'),
        ('Platforms', 'Android, iOS, Web'),
        ('Timeline', '2 Weeks MVP'),
        ('Total Budget', '₹99,800')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = key
        summary_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 1. WHAT WE'RE BUILDING
    doc.add_heading('1. WHAT WE\'RE BUILDING', 1)
    
    doc.add_heading('Three Core Modules', 2)
    modules_table = doc.add_table(rows=4, cols=4)
    modules_table.style = 'Table Grid'
    modules_data = [
        ('Module', 'Features', 'Target', 'Revenue'),
        ('Training Plans', 'Coach profiles, Video sessions, Progress tracking', 'Athletes, Coaches', 'B2B & B2C'),
        ('Events', 'Event listing, Registration, QR check-in', 'Event organizers', 'Commission'),
        ('Ticketing', 'Seat selection, QR tickets, Payment gateway', 'Event attendees', 'Service fee')
    ]
    for i, row_data in enumerate(modules_data):
        for j, cell_data in enumerate(row_data):
            modules_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 2. APP SCREENS
    doc.add_heading('2. APP SCREENS (28 Total)', 1)
    
    doc.add_heading('Navigation Structure', 2)
    nav_items = [
        'Bottom Navigation: Home, Training, Events, Profile',
        'Top Bar: Search, Notifications, Cart',
        'Drawer: Settings, Help, About'
    ]
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Screen Flow', 2)
    flow_items = [
        'Onboarding → Login → Home',
        'Home → Training → Details → Book → Payment → Success',
        'Home → Events → Details → Register → Payment → Success',
        'Home → Tickets → Details → Seat Selection → Payment → Success'
    ]
    for item in flow_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. DETAILED SCREENS
    doc.add_heading('3. DETAILED SCREENS', 1)
    
    # Onboarding
    doc.add_heading('3.1 Onboarding (4 Screens)', 2)
    onboarding_table = doc.add_table(rows=5, cols=3)
    onboarding_table.style = 'Table Grid'
    onboarding_data = [
        ('Screen', 'Elements', 'Purpose'),
        ('Splash', 'Logo, Tagline', 'App loading'),
        ('Welcome 1', 'Image, Title, Next button', 'Find Coaches'),
        ('Welcome 2', 'Image, Title, Next button', 'Discover Events'),
        ('Welcome 3', 'Image, Title, Get Started button', 'Book Tickets')
    ]
    for i, row_data in enumerate(onboarding_data):
        for j, cell_data in enumerate(row_data):
            onboarding_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Authentication
    doc.add_heading('3.2 Authentication (4 Screens)', 2)
    auth_table = doc.add_table(rows=5, cols=3)
    auth_table.style = 'Table Grid'
    auth_data = [
        ('Screen', 'Elements', 'Purpose'),
        ('Login', 'Email, Password, Google/Apple login', 'User login'),
        ('Sign Up', 'Name, Email, Phone, Password, User Type', 'Create account'),
        ('Forgot Password', 'Email input, Reset button', 'Reset password'),
        ('OTP Verify', '6-digit code, Resend OTP', 'Verify phone')
    ]
    for i, row_data in enumerate(auth_data):
        for j, cell_data in enumerate(row_data):
            auth_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # Home
    doc.add_heading('3.3 Home Screen (1 Screen)', 2)
    home_table = doc.add_table(rows=6, cols=2)
    home_table.style = 'Table Grid'
    home_data = [
        ('Section', 'Elements'),
        ('Header', 'Profile picture, Greeting, Search, Notifications, Cart'),
        ('Quick Actions', 'Training Plans, Events, Tickets, Progress (4 cards)'),
        ('Featured', 'Featured Events, Popular Coaches, Trending Plans'),
        ('Upcoming', 'Event cards with Book Now button'),
        ('Recommended', 'Training plan cards with Start Plan button')
    ]
    for i, (key, value) in enumerate(home_data):
        home_table.cell(i, 0).text = key
        home_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # Training
    doc.add_heading('3.4 Training Plans (4 Screens)', 2)
    training_table = doc.add_table(rows=5, cols=2)
    training_table.style = 'Table Grid'
    training_data = [
        ('Screen', 'Elements'),
        ('Plans List', 'Search, Filter, Sort, Plan cards'),
        ('Plan Details', 'Hero image, Coach info, Duration, Level, Price, Schedule, Start Plan button'),
        ('Session', 'Video player, Timer, Exercises list, Complete/Skip buttons'),
        ('Progress', 'Total sessions, Streak, Hours trained, Charts, History')
    ]
    for i, (key, value) in enumerate(training_data):
        training_table.cell(i, 0).text = key
        training_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Events
    doc.add_heading('3.5 Events (4 Screens)', 2)
    events_table = doc.add_table(rows=5, cols=2)
    events_table.style = 'Table Grid'
    events_data = [
        ('Screen', 'Elements'),
        ('Events List', 'Search, Filter, Sort, Event cards'),
        ('Event Details', 'Hero image, Date/Time, Location map, Organizer, Rules, Register button'),
        ('Registration', 'Event summary, Fee breakdown, Participant details, Pay Now button'),
        ('Check-in', 'QR code, Event info, Check-in status')
    ]
    for i, (key, value) in enumerate(events_data):
        events_table.cell(i, 0).text = key
        events_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Tickets
    doc.add_heading('3.6 Tickets (4 Screens)', 2)
    tickets_table = doc.add_table(rows=5, cols=2)
    tickets_table.style = 'Table Grid'
    tickets_data = [
        ('Screen', 'Elements'),
        ('Tickets List', 'Search, Filter, Sort, Ticket cards'),
        ('Ticket Details', 'Event image, Venue, Seat selection, Ticket type, Quantity, Buy Tickets button'),
        ('Seat Selection', 'Interactive seat map, Legend, Selected seats, Price summary'),
        ('Confirmation', 'Booking ID, Event details, Seat details, QR code, Download button')
    ]
    for i, (key, value) in enumerate(tickets_data):
        tickets_table.cell(i, 0).text = key
        tickets_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # Cart & Payment
    doc.add_heading('3.7 Cart & Payment (4 Screens)', 2)
    cart_table = doc.add_table(rows=5, cols=2)
    cart_table.style = 'Table Grid'
    cart_data = [
        ('Screen', 'Elements'),
        ('Cart', 'Cart items, Remove button, Quantity, Promo code, Total, Checkout button'),
        ('Checkout', 'Order summary, Payment methods (Razorpay, UPI, Card), Pay Now button'),
        ('Success', 'Success icon, Amount paid, Transaction ID, View Tickets button'),
        ('Failed', 'Error icon, Error message, Retry button')
    ]
    for i, (key, value) in enumerate(cart_data):
        cart_table.cell(i, 0).text = key
        cart_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Profile
    doc.add_heading('3.8 Profile (4 Screens)', 2)
    profile_table = doc.add_table(rows=5, cols=2)
    profile_table.style = 'Table Grid'
    profile_data = [
        ('Screen', 'Elements'),
        ('Profile', 'Avatar, Name, Email, Phone, Edit Profile, My Bookings/Tickets/Training, Settings, Logout'),
        ('Edit Profile', 'Change avatar, Name, Email, Phone, DOB, Gender, Bio, Sport, Save button'),
        ('My Bookings', 'Tabs (Upcoming/Past), Booking cards, View Details/Cancel buttons'),
        ('My Tickets', 'Tabs (Upcoming/Past), Ticket cards, QR Code/Download buttons')
    ]
    for i, (key, value) in enumerate(profile_data):
        profile_table.cell(i, 0).text = key
        profile_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 4. TECHNICAL ARCHITECTURE
    doc.add_heading('4. TECHNICAL ARCHITECTURE', 1)
    
    doc.add_heading('Tech Stack', 2)
    tech_table = doc.add_table(rows=8, cols=2)
    tech_table.style = 'Table Grid'
    tech_data = [
        ('Layer', 'Technology'),
        ('Frontend', 'Flutter 3.41.9, Dart 3.11.5'),
        ('State Management', 'BLoC'),
        ('Navigation', 'GoRouter'),
        ('HTTP Client', 'Dio'),
        ('Local Storage', 'Hive'),
        ('Backend', 'Firebase / AWS'),
        ('Payment', 'Razorpay')
    ]
    for i, (key, value) in enumerate(tech_data):
        tech_table.cell(i, 0).text = key
        tech_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 5. DEVELOPMENT TIMELINE
    doc.add_heading('5. DEVELOPMENT TIMELINE', 1)
    
    doc.add_heading('2-Week Sprint Plan', 2)
    timeline_table = doc.add_table(rows=11, cols=3)
    timeline_table.style = 'Table Grid'
    timeline_data = [
        ('Day', 'Focus', 'Deliverables'),
        ('1', 'Backend Setup', 'Database, Auth APIs, Basic structure'),
        ('2', 'Flutter Setup', 'Project structure, Theme, Navigation'),
        ('3', 'Auth Screens', 'Login, Signup, OTP, Forgot Password'),
        ('4', 'Home Dashboard', 'Home screen, Quick actions, Featured'),
        ('5', 'Training Plans', 'Plans list, Plan details, Sessions'),
        ('6', 'Events', 'Events list, Event details, Registration'),
        ('7', 'Ticketing', 'Tickets list, Seat selection, Booking'),
        ('8', 'Payments', 'Razorpay integration, Cart, Checkout'),
        ('9', 'Testing', 'Bug fixes, Performance, Polish'),
        ('10', 'Deployment', 'Build, Upload, Submit for review')
    ]
    for i, row_data in enumerate(timeline_data):
        for j, cell_data in enumerate(row_data):
            timeline_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 6. COST BREAKDOWN
    doc.add_heading('6. COST BREAKDOWN', 1)
    
    doc.add_heading('Development Cost', 2)
    dev_table = doc.add_table(rows=5, cols=2)
    dev_table.style = 'Table Grid'
    dev_data = [
        ('Item', 'Cost'),
        ('2 Flutter Developers (10 days)', '₹40,000'),
        ('1 Backend Developer (5 days)', '₹15,000'),
        ('UI/UX Design', '₹10,000'),
        ('Total Development', '₹65,000')
    ]
    for i, (key, value) in enumerate(dev_data):
        dev_table.cell(i, 0).text = key
        dev_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Infrastructure Cost', 2)
    infra_table = doc.add_table(rows=3, cols=2)
    infra_table.style = 'Table Grid'
    infra_data = [
        ('Item', 'Cost'),
        ('AWS (Free Tier)', '₹0'),
        ('Domain', '₹800')
    ]
    for i, (key, value) in enumerate(infra_data):
        infra_table.cell(i, 0).text = key
        infra_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Third-Party Services', 2)
    services_table = doc.add_table(rows=4, cols=2)
    services_table.style = 'Table Grid'
    services_data = [
        ('Item', 'Cost'),
        ('Firebase', '₹0 (Free Tier)'),
        ('Razorpay', '2% per transaction'),
        ('Twilio (OTP)', '₹500')
    ]
    for i, (key, value) in enumerate(services_data):
        services_table.cell(i, 0).text = key
        services_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('App Store Fees', 2)
    store_table = doc.add_table(rows=3, cols=2)
    store_table.style = 'Table Grid'
    store_data = [
        ('Item', 'Cost'),
        ('Google Play Store', '₹18,000 (one-time)'),
        ('Apple Developer', '₹7,500/year')
    ]
    for i, (key, value) in enumerate(store_data):
        store_table.cell(i, 0).text = key
        store_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('TOTAL INVESTMENT', 2)
    total_table = doc.add_table(rows=6, cols=2)
    total_table.style = 'Table Grid'
    total_data = [
        ('Category', 'Amount'),
        ('Development', '₹65,000'),
        ('Infrastructure', '₹800'),
        ('Services', '₹500'),
        ('App Stores', '₹25,500'),
        ('GRAND TOTAL', '₹91,800')
    ]
    for i, (key, value) in enumerate(total_data):
        total_table.cell(i, 0).text = key
        total_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 7. TEAM REQUIREMENTS
    doc.add_heading('7. TEAM REQUIREMENTS', 1)
    
    doc.add_heading('Team Structure', 2)
    team_table = doc.add_table(rows=4, cols=3)
    team_table.style = 'Table Grid'
    team_data = [
        ('Role', 'Count', 'Skills'),
        ('Flutter Developer', '2', 'Flutter, Dart, BLoC'),
        ('Backend Developer', '1', 'Node.js/Firebase, PostgreSQL'),
        ('UI/UX Designer', '1', 'Figma, Design systems')
    ]
    for i, row_data in enumerate(team_data):
        for j, cell_data in enumerate(row_data):
            team_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 8. APP STORE PUBLISHING
    doc.add_heading('8. APP STORE PUBLISHING', 1)
    
    doc.add_heading('Google Play Store', 2)
    google_items = [
        'Create Account: 1-2 days',
        'App Listing: 1 day',
        'Review: 1-7 days',
        'Total: 3-10 days',
        'Cost: ₹18,000 (one-time)'
    ]
    for item in google_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Apple App Store', 2)
    apple_items = [
        'Create Account: 1-2 days',
        'Get D-U-N-S Number: 7-14 days',
        'App Listing: 1 day',
        'Review: 1-2 days',
        'Total: 10-20 days',
        'Cost: ₹7,500/year'
    ]
    for item in apple_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Recommendation', 2)
    rec_items = [
        'Start with Android Only (₹18,000)',
        'Add iOS later after product-market fit',
        '95% Android users in India'
    ]
    for item in rec_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 9. NEXT STEPS
    doc.add_heading('9. NEXT STEPS', 1)
    
    doc.add_heading('Immediate Actions', 2)
    next_table = doc.add_table(rows=6, cols=3)
    next_table.style = 'Table Grid'
    next_data = [
        ('#', 'Action', 'Owner'),
        ('1', 'Review this document', 'Stakeholder'),
        ('2', 'Approve budget', 'Stakeholder'),
        ('3', 'Finalize team', 'Stakeholder'),
        ('4', 'Create GitHub repo', 'Developer'),
        ('5', 'Start Day 1', 'Team')
    ]
    for i, row_data in enumerate(next_data):
        for j, cell_data in enumerate(row_data):
            next_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 10. CONCLUSION
    doc.add_heading('10. CONCLUSION', 1)
    
    conclusion_table = doc.add_table(rows=7, cols=2)
    conclusion_table.style = 'Table Grid'
    conclusion_data = [
        ('Item', 'Status'),
        ('Documentation', 'Complete'),
        ('UI/UX Specification', 'Complete'),
        ('Cost Estimation', 'Complete'),
        ('Timeline', '2 weeks'),
        ('Team', 'To be finalized'),
        ('Development', 'Ready to start')
    ]
    for i, (key, value) in enumerate(conclusion_data):
        conclusion_table.cell(i, 0).text = key
        conclusion_table.cell(i, 1).text = value
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_SHOWCASE.docx')
    doc.save(output_path)
    print(f"Showcase document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_showcase()
