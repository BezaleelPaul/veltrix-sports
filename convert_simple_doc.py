"""
Convert simple requirements document to Word
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_simple_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Requirements & Cost Summary', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.add_run('Date: ').bold = True
    info.add_run('August 29, 2026\n')
    info.add_run('For: ').bold = True
    info.add_run('Stakeholder Review')
    
    doc.add_page_break()
    
    # 1. PROJECT REQUIREMENTS
    doc.add_heading('1. PROJECT REQUIREMENTS', 1)
    
    doc.add_heading('1.1 What We\'re Building', 2)
    doc.add_paragraph('A sports platform with 3 core modules:')
    
    module_table = doc.add_table(rows=4, cols=2)
    module_table.style = 'Table Grid'
    module_data = [
        ('Module', 'Description'),
        ('Training Plans', 'Create, manage, and follow workout plans'),
        ('Events', 'Browse, register, and manage sports events'),
        ('Ticketing', 'Buy, sell, and manage event tickets')
    ]
    for i, (key, value) in enumerate(module_data):
        module_table.cell(i, 0).text = key
        module_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('1.2 Platform Requirements', 2)
    platform_table = doc.add_table(rows=4, cols=2)
    platform_table.style = 'Table Grid'
    platform_data = [
        ('Platform', 'Required'),
        ('Android', 'Yes'),
        ('iOS', 'Yes'),
        ('Web', 'Yes')
    ]
    for i, (key, value) in enumerate(platform_data):
        platform_table.cell(i, 0).text = key
        platform_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('1.3 Feature Requirements', 2)
    
    doc.add_heading('Training Plans (Must Have)', 3)
    training_features = [
        'User registration & login',
        'Create training plans',
        'AI-generated plans',
        'Track workout sessions',
        'View progress charts',
        'Connect fitness devices'
    ]
    for feature in training_features:
        doc.add_paragraph('☐ ' + feature)
    
    doc.add_heading('Events (Must Have)', 3)
    event_features = [
        'Browse events',
        'Filter by sport/city/date',
        'Register for events',
        'View my registrations',
        'Event details with route map'
    ]
    for feature in event_features:
        doc.add_paragraph('☐ ' + feature)
    
    doc.add_heading('Ticketing (Should Have)', 3)
    ticket_features = [
        'Buy tickets',
        'View purchased tickets',
        'QR code for check-in',
        'Promo code support'
    ]
    for feature in ticket_features:
        doc.add_paragraph('☐ ' + feature)
    
    doc.add_heading('Payments (Must Have)', 3)
    payment_features = [
        'Razorpay integration',
        'UPI, Cards, Net Banking',
        'Payment receipts',
        'Refund support'
    ]
    for feature in payment_features:
        doc.add_paragraph('☐ ' + feature)
    
    doc.add_page_break()
    
    # 2. TOOLS & TECHNOLOGY
    doc.add_heading('2. TOOLS & TECHNOLOGY', 1)
    
    doc.add_heading('2.1 Development Tools (FREE)', 2)
    tools_table = doc.add_table(rows=7, cols=3)
    tools_table.style = 'Table Grid'
    tools_data = [
        ('Tool', 'Cost', 'Purpose'),
        ('Flutter SDK', '₹0', 'Mobile app development'),
        ('Dart SDK', '₹0', 'Programming language'),
        ('VS Code', '₹0', 'Code editor'),
        ('Git', '₹0', 'Version control'),
        ('Node.js', '₹0', 'Backend runtime'),
        ('Total', '₹0', '')
    ]
    for i, row_data in enumerate(tools_data):
        for j, cell_data in enumerate(row_data):
            tools_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('2.2 Cloud Services (AWS)', 2)
    aws_table = doc.add_table(rows=6, cols=3)
    aws_table.style = 'Table Grid'
    aws_data = [
        ('Service', 'Free Tier', 'After Free Tier'),
        ('EC2 (Compute)', '750 hrs/month', '₹1,400/month'),
        ('RDS (Database)', '750 hrs/month', '₹2,800/month'),
        ('S3 (Storage)', '5GB', '₹200/month'),
        ('CloudFront (CDN)', '1TB transfer', '₹800/month'),
        ('Total', '₹0 (12 months)', '₹6,600/month')
    ]
    for i, row_data in enumerate(aws_data):
        for j, cell_data in enumerate(row_data):
            aws_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('2.3 Third-Party Services', 2)
    third_table = doc.add_table(rows=5, cols=3)
    third_table.style = 'Table Grid'
    third_data = [
        ('Service', 'Cost', 'Purpose'),
        ('Razorpay', '2% per transaction', 'Payments'),
        ('Firebase', 'Free (basic)', 'Push notifications'),
        ('Garmin API', 'Free', 'Device integration'),
        ('Total', '~₹500/month', '')
    ]
    for i, row_data in enumerate(third_data):
        for j, cell_data in enumerate(row_data):
            third_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 3. COST BREAKDOWN
    doc.add_heading('3. COST BREAKDOWN', 1)
    
    # Option A
    doc.add_heading('3.1 Option A: MVP Only (2 Weeks)', 2)
    doc.add_paragraph('What You Get: Basic login/register, Dashboard, View training plans, View events, Basic payment')
    
    mvp_table = doc.add_table(rows=8, cols=2)
    mvp_table.style = 'Table Grid'
    mvp_data = [
        ('Item', 'Cost'),
        ('Flutter Developer (2)', '₹20,000'),
        ('Backend Developer (1)', '₹10,000'),
        ('UI/UX Designer (1)', '₹7,500'),
        ('QA Engineer (1)', '₹6,250'),
        ('AWS (Free Tier)', '₹0'),
        ('Third-party services', '₹1,000'),
        ('TOTAL', '₹44,750')
    ]
    for i, (key, value) in enumerate(mvp_data):
        mvp_table.cell(i, 0).text = key
        mvp_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Option B
    doc.add_heading('3.2 Option B: Full Product (3 Months)', 2)
    doc.add_paragraph('What You Get: Complete training plan module, Complete events module, Complete ticketing module, Payment integration, Device integration, Admin panel')
    
    full_table = doc.add_table(rows=9, cols=2)
    full_table.style = 'Table Grid'
    full_data = [
        ('Item', 'Cost'),
        ('Flutter Developer (3)', '₹1,80,000'),
        ('Backend Developer (2)', '₹1,20,000'),
        ('UI/UX Designer (1)', '₹45,000'),
        ('QA Engineer (2)', '₹75,000'),
        ('DevOps (1)', '₹75,000'),
        ('AWS (3 months)', '₹19,800'),
        ('Third-party services', '₹1,500'),
        ('TOTAL', '₹5,16,300')
    ]
    for i, (key, value) in enumerate(full_data):
        full_table.cell(i, 0).text = key
        full_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    # Option C
    doc.add_heading('3.3 Option C: Enterprise (6 Months)', 2)
    doc.add_paragraph('What You Get: Everything in Full Product + AI plan generation, Live event tracking, Advanced analytics, Multi-language support, White-label options')
    
    ent_table = doc.add_table(rows=10, cols=2)
    ent_table.style = 'Table Grid'
    ent_data = [
        ('Item', 'Cost'),
        ('Flutter Developer (4)', '₹4,80,000'),
        ('Backend Developer (3)', '₹3,60,000'),
        ('UI/UX Designer (2)', '₹1,20,000'),
        ('QA Engineer (2)', '₹1,50,000'),
        ('DevOps (1)', '₹1,20,000'),
        ('Project Manager (1)', '₹60,000'),
        ('AWS (6 months)', '₹39,600'),
        ('Third-party services', '₹3,000'),
        ('TOTAL', '₹13,32,600')
    ]
    for i, (key, value) in enumerate(ent_data):
        ent_table.cell(i, 0).text = key
        ent_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 4. COMPARISON TABLE
    doc.add_heading('4. COMPARISON TABLE', 1)
    
    comp_table = doc.add_table(rows=12, cols=4)
    comp_table.style = 'Table Grid'
    comp_data = [
        ('Feature', 'MVP (₹44K)', 'Full (₹5.1L)', 'Enterprise (₹13.3L)'),
        ('Login/Register', 'Yes', 'Yes', 'Yes'),
        ('Dashboard', 'Yes', 'Yes', 'Yes'),
        ('Training Plans', 'Basic', 'Full', 'Full + AI'),
        ('Events', 'Basic', 'Full', 'Full + Live'),
        ('Ticketing', 'No', 'Yes', 'Yes'),
        ('Payments', 'Basic', 'Full', 'Full'),
        ('Device Sync', 'No', 'Yes', 'Yes'),
        ('Admin Panel', 'No', 'Yes', 'Yes'),
        ('Analytics', 'No', 'Basic', 'Advanced'),
        ('Multi-language', 'No', 'No', 'Yes'),
        ('Timeline', '2 weeks', '3 months', '6 months')
    ]
    for i, row_data in enumerate(comp_data):
        for j, cell_data in enumerate(row_data):
            comp_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 5. RECOMMENDATION
    doc.add_heading('5. RECOMMENDATION', 1)
    
    doc.add_heading('For Startup/MVP', 2)
    doc.add_paragraph('Choose Option A (₹44,750)')
    doc.add_paragraph('• Get basic product running')
    doc.add_paragraph('• Test with users')
    doc.add_paragraph('• Iterate based on feedback')
    
    doc.add_heading('For Business', 2)
    doc.add_paragraph('Choose Option B (₹5,16,300)')
    doc.add_paragraph('• Complete product')
    doc.add_paragraph('• All features working')
    doc.add_paragraph('• Ready for market')
    
    doc.add_heading('For Enterprise', 2)
    doc.add_paragraph('Choose Option C (₹13,32,600)')
    doc.add_paragraph('• Full-featured platform')
    doc.add_paragraph('• Scalable architecture')
    doc.add_paragraph('• Advanced capabilities')
    
    doc.add_page_break()
    
    # 6. TIMELINE
    doc.add_heading('6. TIMELINE', 1)
    
    doc.add_heading('MVP (2 Weeks)', 2)
    mvp_timeline = [
        'Week 1: Day 1-2: Setup + Auth, Day 3-5: Dashboard + Plans, Day 6-7: Events + Basic UI',
        'Week 2: Day 8-9: Integration + Testing, Day 10: Deploy'
    ]
    for item in mvp_timeline:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Full Product (3 Months)', 2)
    full_timeline = [
        'Month 1: Auth System, Dashboard, Core Infrastructure',
        'Month 2: Training Plans, Events, Ticketing',
        'Month 3: Payments, Device Sync, Testing + Deploy'
    ]
    for item in full_timeline:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. NEXT STEPS
    doc.add_heading('7. NEXT STEPS', 1)
    
    steps_table = doc.add_table(rows=6, cols=3)
    steps_table.style = 'Table Grid'
    steps_data = [
        ('Step', 'Action', 'Timeline'),
        ('1', 'Approve requirements', 'Today'),
        ('2', 'Select option (A/B/C)', 'Today'),
        ('3', 'Finalize team', 'Day 1'),
        ('4', 'Start development', 'Day 2'),
        ('5', 'MVP delivery', 'Week 2')
    ]
    for i, row_data in enumerate(steps_data):
        for j, cell_data in enumerate(row_data):
            steps_table.cell(i, j).text = cell_data
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_REQUIREMENTS.docx')
    doc.save(output_path)
    print(f"Simple requirements document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_simple_document()
