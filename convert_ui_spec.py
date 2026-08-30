"""
Convert UI/UX specification to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_ui_spec():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('UI/UX Specification - Mobile App Screens & Features', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. NAVIGATION STRUCTURE
    doc.add_heading('1. APP NAVIGATION STRUCTURE', 1)
    
    doc.add_heading('1.1 Bottom Navigation Bar', 2)
    nav_table = doc.add_table(rows=5, cols=3)
    nav_table.style = 'Table Grid'
    nav_data = [
        ('Icon', 'Label', 'Screen'),
        ('Home', 'Home', 'Home Screen'),
        ('Training', 'Training', 'Training Plans'),
        ('Events', 'Events', 'Events List'),
        ('Profile', 'Profile', 'User Profile')
    ]
    for i, row_data in enumerate(nav_data):
        for j, cell_data in enumerate(row_data):
            nav_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    # 2. ONBOARDING SCREENS
    doc.add_heading('2. ONBOARDING SCREENS', 1)
    
    screens = [
        ('Splash Screen', [
            'Logo: Veltrix Sports logo centered',
            'Tagline: "Train. Compete. Win."',
            'Duration: 2-3 seconds',
            'Animation: Fade in logo'
        ]),
        ('Welcome Screen 1', [
            'Image: Athlete training illustration',
            'Title: "Find Expert Coaches"',
            'Description: "Connect with certified coaches"',
            'Button: "Next"',
            'Dots: Page indicator (1/3)'
        ]),
        ('Welcome Screen 2', [
            'Image: Sports event illustration',
            'Title: "Discover Events"',
            'Description: "Find and register for sports events"',
            'Button: "Next"',
            'Dots: Page indicator (2/3)'
        ]),
        ('Welcome Screen 3', [
            'Image: Ticket booking illustration',
            'Title: "Book Tickets"',
            'Description: "Get tickets for tournaments"',
            'Button: "Get Started"',
            'Dots: Page indicator (3/3)'
        ])
    ]
    
    for screen_name, elements in screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. AUTHENTICATION SCREENS
    doc.add_heading('3. AUTHENTICATION SCREENS', 1)
    
    auth_screens = [
        ('Login Screen', [
            'Email: Input field',
            'Password: Input with show/hide',
            'Remember Me: Checkbox',
            'Login: Primary button',
            'Forgot Password: Link',
            'Google: Sign in button',
            'Apple: Sign in button',
            'Create Account: Link'
        ]),
        ('Sign Up Screen', [
            'Full Name: Input field',
            'Email: Input field',
            'Phone: Input field',
            'Password: Input field',
            'Confirm Password: Input field',
            'I\'m a: Selection (Coach/Athlete/Both)',
            'Create Account: Primary button',
            'Terms: Link',
            'Already have account: Link'
        ]),
        ('Forgot Password Screen', [
            'Email: Input field',
            'Send Reset Link: Primary button',
            'Back to Login: Link'
        ]),
        ('OTP Verification Screen', [
            'OTP Input: 6-digit code',
            'Resend OTP: Link',
            'Verify: Primary button',
            'Timer: Countdown'
        ])
    ]
    
    for screen_name, elements in auth_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. HOME SCREEN
    doc.add_heading('4. HOME SCREEN', 1)
    
    home_sections = [
        ('Header Section', [
            'Profile Picture: Avatar',
            'Greeting: "Good Morning, [Name]"',
            'Search: Icon button',
            'Notifications: Icon button',
            'Cart: Icon button'
        ]),
        ('Quick Actions Section', [
            'Training Plans: Card with icon',
            'Events: Card with icon',
            'Tickets: Card with icon',
            'Progress: Card with icon'
        ]),
        ('Featured Section', [
            'Featured Events: Horizontal scroll cards',
            'Popular Coaches: Horizontal scroll cards',
            'Trending Plans: Horizontal scroll cards',
            'See All: Link'
        ]),
        ('Upcoming Events', [
            'Event Card: Image, name, date, location, price',
            'Book Now: Button',
            'View Details: Button',
            'See All: Link'
        ]),
        ('Recommended Training', [
            'Plan Card: Image, name, coach, rating, price',
            'Start Plan: Button',
            'View Details: Button',
            'See All: Link'
        ])
    ]
    
    for section_name, elements in home_sections:
        doc.add_heading(section_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. TRAINING PLANS
    doc.add_heading('5. TRAINING PLANS SCREENS', 1)
    
    training_screens = [
        ('Training Plans List', [
            'Search Bar: Search training plans',
            'Filter: Filter by sport, level, price',
            'Sort: Sort by rating, price, popular',
            'Plan Card: Image, name, coach, rating, price',
            'View Details: Button'
        ]),
        ('Training Plan Details', [
            'Hero Image: Plan banner',
            'Plan Name: Title',
            'Coach Info: Avatar, name, rating',
            'Duration: Number of weeks',
            'Level: Beginner/Intermediate/Advanced',
            'Price: Amount',
            'Description: Plan overview',
            'Schedule: Weekly breakdown',
            'Reviews: User reviews',
            'Start Plan: Button',
            'Add to Cart: Button'
        ]),
        ('Training Session Screen', [
            'Session Title: Current session name',
            'Video Player: Training video',
            'Timer: Workout timer',
            'Exercises: List of exercises',
            'Reps/Sets: Exercise details',
            'Complete: Button',
            'Skip: Button',
            'Next: Button'
        ]),
        ('Progress Dashboard', [
            'Total Sessions: Number completed',
            'Streak: Days in a row',
            'Hours Trained: Total time',
            'Charts: Weekly progress graph',
            'History: Past sessions'
        ])
    ]
    
    for screen_name, elements in training_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. EVENTS
    doc.add_heading('6. EVENTS SCREENS', 1)
    
    event_screens = [
        ('Events List', [
            'Search Bar: Search events',
            'Filter: Filter by sport, location, date',
            'Sort: Sort by date, price, popular',
            'Event Card: Image, name, date, location, price',
            'View Details: Button',
            'Map View: Toggle map view'
        ]),
        ('Event Details', [
            'Hero Image: Event banner',
            'Event Name: Title',
            'Date & Time: When event happens',
            'Location: Venue with map',
            'Organizer: Event organizer info',
            'Description: Event details',
            'Rules: Event rules',
            'Prize Pool: Prize information',
            'Participants: Registered count',
            'Register: Button',
            'Add to Calendar: Button',
            'Share: Button'
        ]),
        ('Event Registration', [
            'Event Summary: Name, date, location',
            'Registration Fee: Amount breakdown',
            'Participant Details: Name, age, gender, phone',
            'Emergency Contact: Name, phone',
            'Medical Info: Any health conditions',
            'Terms: Accept terms checkbox',
            'Pay Now: Button',
            'Pay Later: Button'
        ]),
        ('Event Check-in', [
            'QR Code: Unique check-in code',
            'Event Info: Name, date',
            'Check-in Status: Verified/Pending',
            'Share QR: Button'
        ])
    ]
    
    for screen_name, elements in event_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. TICKET BOOKING
    doc.add_heading('7. TICKET BOOKING SCREENS', 1)
    
    ticket_screens = [
        ('Tickets List', [
            'Search Bar: Search tickets',
            'Filter: Filter by event type, date, price',
            'Sort: Sort by date, price',
            'Ticket Card: Event name, date, venue, price',
            'Buy Now: Button'
        ]),
        ('Ticket Details', [
            'Event Image: Event banner',
            'Event Name: Title',
            'Date & Time: When event happens',
            'Venue: Location with map',
            'Seat Selection: Interactive seat map',
            'Ticket Type: VIP/General/Student',
            'Price: Per ticket price',
            'Quantity: +/- selector',
            'Total: Total amount',
            'Buy Tickets: Button'
        ]),
        ('Seat Selection', [
            'Seat Map: Interactive venue map',
            'Legend: Available/Selected/Sold',
            'Selected Seats: List of selected',
            'Price Summary: Total calculation',
            'Proceed: Button'
        ]),
        ('Booking Confirmation', [
            'Booking ID: Unique identifier',
            'Event Details: Name, date, venue',
            'Seat Details: Section, row, seat',
            'QR Code: Ticket QR code',
            'Download: Button',
            'Add to Wallet: Button',
            'Share: Button'
        ])
    ]
    
    for screen_name, elements in ticket_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 8. CART & PAYMENT
    doc.add_heading('8. CART & PAYMENT SCREENS', 1)
    
    cart_screens = [
        ('Cart Screen', [
            'Cart Items: List of items',
            'Remove: Button per item',
            'Quantity: +/- per item',
            'Subtotal: Item total',
            'Discount: Promo code input',
            'Apply: Button',
            'Total: Final amount',
            'Checkout: Button'
        ]),
        ('Checkout Screen', [
            'Order Summary: Items list',
            'Delivery Info: Email for tickets',
            'Payment Methods: Select payment',
            'Razorpay: Button',
            'UPI: Button',
            'Card: Button',
            'Net Banking: Button',
            'Pay Now: Button'
        ]),
        ('Payment Success', [
            'Success Icon: Green checkmark',
            'Amount Paid: Total amount',
            'Transaction ID: Unique ID',
            'Booking ID: Unique ID',
            'Email Sent: Confirmation email',
            'View Tickets: Button',
            'Back to Home: Button'
        ]),
        ('Payment Failed', [
            'Error Icon: Red cross',
            'Error Message: What went wrong',
            'Retry: Button',
            'Cancel: Button',
            'Support: Link'
        ])
    ]
    
    for screen_name, elements in cart_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 9. USER PROFILE
    doc.add_heading('9. USER PROFILE SCREENS', 1)
    
    profile_screens = [
        ('Profile Screen', [
            'Profile Picture: Avatar with edit',
            'Name: User\'s name',
            'Email: Email address',
            'Phone: Phone number',
            'Member Since: Join date',
            'Edit Profile: Button',
            'My Bookings: Link',
            'My Tickets: Link',
            'My Training: Link',
            'Settings: Link',
            'Help: Link',
            'Logout: Button'
        ]),
        ('Edit Profile Screen', [
            'Profile Picture: Change avatar',
            'Full Name: Edit name',
            'Email: Edit email',
            'Phone: Edit phone',
            'Date of Birth: Date picker',
            'Gender: Male/Female/Other',
            'Bio: Text area',
            'Sport: Select sports',
            'Save: Button'
        ]),
        ('My Bookings', [
            'Tabs: Upcoming / Past',
            'Booking Card: Event name, date, status',
            'View Details: Button',
            'Cancel: Button',
            'Download: Button'
        ]),
        ('My Tickets', [
            'Tabs: Upcoming / Past',
            'Ticket Card: Event name, date, seat',
            'QR Code: Button',
            'Download: Button',
            'Transfer: Button'
        ])
    ]
    
    for screen_name, elements in profile_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. COACH PROFILE
    doc.add_heading('10. COACH PROFILE SCREENS', 1)
    
    coach_screens = [
        ('Coach List', [
            'Search Bar: Search coaches',
            'Filter: Filter by sport, rating, price',
            'Coach Card: Avatar, name, sport, rating, price',
            'View Profile: Button'
        ]),
        ('Coach Profile', [
            'Profile Picture: Large avatar',
            'Name: Coach\'s name',
            'Sport: Specialization',
            'Rating: Star rating',
            'Reviews: Number of reviews',
            'Experience: Years of experience',
            'Bio: About coach',
            'Plans: Training plans offered',
            'Events: Events organized',
            'Book Session: Button',
            'View Plans: Button',
            'Message: Button'
        ]),
        ('Book Session with Coach', [
            'Coach Info: Name, photo',
            'Session Type: 1-on-1 / Group',
            'Date: Date picker',
            'Time: Time slot selection',
            'Duration: 30min / 60min / 90min',
            'Location: Online / In-person',
            'Price: Session cost',
            'Confirm Booking: Button'
        ])
    ]
    
    for screen_name, elements in coach_screens:
        doc.add_heading(screen_name, 2)
        for element in elements:
            doc.add_paragraph(element, style='List Bullet')
    
    doc.add_page_break()
    
    # 11. BUTTONS & ACTIONS
    doc.add_heading('11. BUTTONS & ACTIONS', 1)
    
    doc.add_heading('11.1 Primary Buttons', 2)
    primary_table = doc.add_table(rows=7, cols=3)
    primary_table.style = 'Table Grid'
    primary_data = [
        ('Button', 'Color', 'Usage'),
        ('Login', 'Blue', 'Login action'),
        ('Sign Up', 'Blue', 'Create account'),
        ('Book Now', 'Blue', 'Book event/ticket'),
        ('Start Plan', 'Blue', 'Start training'),
        ('Pay Now', 'Green', 'Payment action'),
        ('Confirm', 'Green', 'Confirm action')
    ]
    for i, row_data in enumerate(primary_data):
        for j, cell_data in enumerate(row_data):
            primary_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('11.2 Secondary Buttons', 2)
    secondary_table = doc.add_table(rows=5, cols=3)
    secondary_table.style = 'Table Grid'
    secondary_data = [
        ('Button', 'Style', 'Usage'),
        ('Cancel', 'Outline', 'Cancel action'),
        ('Skip', 'Text', 'Skip step'),
        ('Back', 'Text', 'Go back'),
        ('View Details', 'Outline', 'View more')
    ]
    for i, row_data in enumerate(secondary_data):
        for j, cell_data in enumerate(row_data):
            secondary_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('11.3 Icon Buttons', 2)
    icon_table = doc.add_table(rows=6, cols=2)
    icon_table.style = 'Table Grid'
    icon_data = [
        ('Icon', 'Action'),
        ('Heart', 'Favorite/Like'),
        ('Share', 'Share'),
        ('Search', 'Search'),
        ('Bell', 'Notifications'),
        ('Cart', 'Cart')
    ]
    for i, row_data in enumerate(icon_data):
        for j, cell_data in enumerate(row_data):
            icon_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 12. TOAST MESSAGES
    doc.add_heading('12. TOAST MESSAGES', 1)
    
    doc.add_heading('12.1 Success', 2)
    success_items = [
        '"Booking Confirmed!" - After booking',
        '"Payment Successful!" - After payment',
        '"Profile Updated!" - After edit',
        '"Added to Cart!" - After add to cart'
    ]
    for item in success_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('12.2 Error', 2)
    error_items = [
        '"Booking Failed" - On error',
        '"Payment Failed" - On error',
        '"Something Went Wrong" - On error',
        '"No Internet Connection" - On offline'
    ]
    for item in error_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('12.3 Warning', 2)
    warning_items = [
        '"Session Expired" - On timeout',
        '"Item Removed" - After delete',
        '"Cart Empty" - On checkout'
    ]
    for item in warning_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 13. SCREEN FLOW
    doc.add_heading('13. SCREEN FLOW DIAGRAM', 1)
    
    flow_text = """
Onboarding → Login/Signup → Home
                              ↓
                    ┌─────────┼─────────┐
                    ↓         ↓         ↓
                Training   Events    Tickets
                    ↓         ↓         ↓
                Details   Details   Details
                    ↓         ↓         ↓
                Book/Start Register  Book
                    ↓         ↓         ↓
                Payment   Payment   Payment
                    ↓         ↓         ↓
                Success   Success   Success
                    ↓         ↓         ↓
                    └─────────┼─────────┘
                              ↓
                           Profile
    """
    
    p = doc.add_paragraph(flow_text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # 14. RESPONSIVE DESIGN
    doc.add_heading('14. RESPONSIVE DESIGN', 1)
    
    doc.add_heading('14.1 Screen Sizes', 2)
    size_table = doc.add_table(rows=7, cols=4)
    size_table.style = 'Table Grid'
    size_data = [
        ('Device', 'Width', 'Height', 'Size'),
        ('iPhone SE', '375', '667', 'Small'),
        ('iPhone 14', '390', '844', 'Medium'),
        ('iPhone 14 Pro Max', '430', '932', 'Large'),
        ('Android Small', '360', '640', 'Small'),
        ('Android Medium', '390', '844', 'Medium'),
        ('Android Large', '412', '915', 'Large')
    ]
    for i, row_data in enumerate(size_data):
        for j, cell_data in enumerate(row_data):
            size_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('14.2 Breakpoints', 2)
    bp_table = doc.add_table(rows=4, cols=3)
    bp_table.style = 'Table Grid'
    bp_data = [
        ('Breakpoint', 'Value', 'Usage'),
        ('Mobile', '< 600px', 'Single column'),
        ('Tablet', '600-900px', 'Two columns'),
        ('Desktop', '> 900px', 'Three columns')
    ]
    for i, row_data in enumerate(bp_data):
        for j, cell_data in enumerate(row_data):
            bp_table.cell(i, j).text = cell_data
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_UI_SPECIFICATION.docx')
    doc.save(output_path)
    print(f"UI specification saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_ui_spec()
