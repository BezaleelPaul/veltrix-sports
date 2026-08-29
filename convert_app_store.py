"""
Convert app store guide to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_app_store_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Google Play & Apple App Store Publishing Guide', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. GOOGLE PLAY STORE
    doc.add_heading('1. GOOGLE PLAY STORE', 1)
    
    doc.add_heading('1.1 Prerequisites', 2)
    prereq_table = doc.add_table(rows=5, cols=2)
    prereq_table.style = 'Table Grid'
    prereq_data = [
        ('Item', 'Requirement'),
        ('Google Account', 'Gmail account'),
        ('Developer Fee', '₹18,000 (One-time)'),
        ('Bank Account', 'For payments'),
        ('GST Number', 'For Indian developers')
    ]
    for i, (key, value) in enumerate(prereq_data):
        prereq_table.cell(i, 0).text = key
        prereq_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('1.2 Step-by-Step Process', 2)
    
    doc.add_heading('Step 1: Create Developer Account', 3)
    step1_items = [
        'Go to https://play.google.com/console',
        'Sign in with Gmail',
        'Click "Create Developer Account"',
        'Pay ₹18,000 (one-time fee)',
        'Fill developer details',
        'Verify email',
        'Account created in 24-48 hours'
    ]
    for item in step1_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 2: Create App', 3)
    step2_items = [
        'Go to Google Play Console',
        'Click "Create App"',
        'Fill app name: "Veltrix Sports"',
        'Select: Free, App',
        'Accept declarations',
        'Click "Create App"'
    ]
    for item in step2_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 3: Complete Store Listing', 3)
    step3_items = [
        'Go to "Store presence" → "Main store listing"',
        'Add app name (max 30 chars)',
        'Add short description (max 80 chars)',
        'Add full description (max 4000 chars)',
        'Upload app icon (512x512 PNG)',
        'Upload feature graphic (1024x500 PNG)',
        'Upload screenshots (min 2, max 8)'
    ]
    for item in step3_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 4: Content Rating', 3)
    step4_items = [
        'Go to "Store presence" → "Content rating"',
        'Fill questionnaire',
        'Get rating (Everyone, Teen, Mature)',
        'Submit for rating'
    ]
    for item in step4_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 5: Pricing & Distribution', 3)
    step5_items = [
        'Go to "Store presence" → "Pricing & distribution"',
        'Set price: Free',
        'Select countries',
        'Accept terms'
    ]
    for item in step5_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 6: Upload APK/AAB', 3)
    step6_items = [
        'Go to "Production" → "Create new release"',
        'Upload AAB file (recommended)',
        'Add release notes',
        'Review and start rollout'
    ]
    for item in step6_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 7: Submit for Review', 3)
    step7_items = [
        'Submit for review',
        'Review takes 1-7 days',
        'Fix any issues',
        'App goes live'
    ]
    for item in step7_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 2. APPLE APP STORE
    doc.add_heading('2. APPLE APP STORE', 1)
    
    doc.add_heading('2.1 Prerequisites', 2)
    apple_prereq = doc.add_table(rows=6, cols=2)
    apple_prereq.style = 'Table Grid'
    apple_data = [
        ('Item', 'Requirement'),
        ('Apple ID', 'Apple account'),
        ('Developer Fee', '₹7,500/year'),
        ('Mac Computer', 'For Xcode'),
        ('Bank Account', 'For payments'),
        ('GST Number', 'For Indian developers')
    ]
    for i, (key, value) in enumerate(apple_data):
        apple_prereq.cell(i, 0).text = key
        apple_prereq.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('2.2 Step-by-Step Process', 2)
    
    doc.add_heading('Step 1: Create Developer Account', 3)
    apple_step1 = [
        'Go to https://developer.apple.com',
        'Sign in with Apple ID',
        'Click "Join the Apple Developer Program"',
        'Pay ₹7,500/year',
        'Fill legal entity details',
        'Verify organization',
        'Account approved in 24-48 hours'
    ]
    for item in apple_step1:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 2: Get D-U-N-S Number', 3)
    apple_step2 = [
        'Go to https://developer.apple.com/enroll/duns-lookup/',
        'Search your company',
        'If not found, request new D-U-N-S number',
        'Wait 7-14 business days',
        'Free service from Apple'
    ]
    for item in apple_step2:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 3: Create App ID', 3)
    apple_step3 = [
        'Go to Apple Developer → Certificates, Identifiers & Profiles',
        'Click "+" to create App ID',
        'Select "App IDs"',
        'Fill description: "Veltrix Sports"',
        'Bundle ID: "com.veltrix.sports"',
        'Enable Push Notifications',
        'Register'
    ]
    for item in apple_step3:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 4: App Store Connect', 3)
    apple_step4 = [
        'Go to https://appstoreconnect.apple.com',
        'Click "My Apps"',
        'Click "+" → "New App"',
        'Platform: iOS',
        'Name: "Veltrix Sports"',
        'Bundle ID: "com.veltrix.sports"'
    ]
    for item in apple_step4:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 5: Complete App Information', 3)
    apple_step5 = [
        'Add app name (max 30 chars)',
        'Add subtitle (max 30 chars)',
        'Add description (max 4000 chars)',
        'Add keywords (max 100 chars)',
        'Upload app icon (1024x1024 PNG)',
        'Upload screenshots for all devices'
    ]
    for item in apple_step5:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 6: Build Upload', 3)
    apple_step6 = [
        'Open project in Xcode',
        'Select "Any iOS Device"',
        'Product → Archive',
        'Validate app',
        'Upload to App Store Connect'
    ]
    for item in apple_step6:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_heading('Step 7: Submit for Review', 3)
    apple_step7 = [
        'Go to "App Store" tab',
        'Click "Submit for Review"',
        'Answer review questions',
        'Submit',
        'Review takes 24-48 hours'
    ]
    for item in apple_step7:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 3. COMMON REQUIREMENTS
    doc.add_heading('3. COMMON REQUIREMENTS', 1)
    
    doc.add_heading('3.1 App Icons', 2)
    icon_table = doc.add_table(rows=3, cols=3)
    icon_table.style = 'Table Grid'
    icon_data = [
        ('Store', 'Size', 'Format'),
        ('Google Play', '512x512', 'PNG'),
        ('Apple App Store', '1024x1024', 'PNG')
    ]
    for i, row_data in enumerate(icon_data):
        for j, cell_data in enumerate(row_data):
            icon_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('')
    
    doc.add_heading('3.2 Screenshots', 2)
    screenshot_items = [
        'Google Play: Phone min 2, max 8',
        'Apple: iPhone + iPad required',
        'Ratio: 16:9 or 9:16'
    ]
    for item in screenshot_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.3 Privacy Policy', 2)
    privacy_items = [
        'What data you collect',
        'How you use data',
        'Data sharing practices',
        'User rights',
        'Contact information'
    ]
    for item in privacy_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. TIMELINE
    doc.add_heading('4. TIMELINE', 1)
    
    doc.add_heading('Google Play Store', 2)
    google_timeline = doc.add_table(rows=5, cols=2)
    google_timeline.style = 'Table Grid'
    google_data = [
        ('Step', 'Duration'),
        ('Account creation', '1-2 days'),
        ('App listing', '1 day'),
        ('Review', '1-7 days'),
        ('Total', '3-10 days')
    ]
    for i, (key, value) in enumerate(google_data):
        google_timeline.cell(i, 0).text = key
        google_timeline.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Apple App Store', 2)
    apple_timeline = doc.add_table(rows=6, cols=2)
    apple_timeline.style = 'Table Grid'
    apple_data = [
        ('Step', 'Duration'),
        ('Account creation', '1-2 days'),
        ('D-U-N-S Number', '7-14 days'),
        ('App listing', '1 day'),
        ('Review', '1-2 days'),
        ('Total', '10-20 days')
    ]
    for i, (key, value) in enumerate(apple_data):
        apple_timeline.cell(i, 0).text = key
        apple_timeline.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 5. COSTS
    doc.add_heading('5. COSTS SUMMARY', 1)
    
    cost_table = doc.add_table(rows=4, cols=2)
    cost_table.style = 'Table Grid'
    cost_data = [
        ('Store', 'Cost'),
        ('Google Play', '₹18,000'),
        ('Apple App Store', '₹7,500/year'),
        ('Total (Both)', '₹25,500 first year')
    ]
    for i, (key, value) in enumerate(cost_data):
        cost_table.cell(i, 0).text = key
        cost_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 6. CHECKLIST
    doc.add_heading('6. CHECKLIST', 1)
    
    doc.add_heading('Before Publishing', 2)
    before_items = [
        'App tested on real devices',
        'No bugs or crashes',
        'Privacy policy ready',
        'Screenshots taken',
        'App description written',
        'App icon ready',
        'Store listing complete'
    ]
    for item in before_items:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Google Play', 2)
    google_items = [
        'Developer account created',
        'App created in console',
        'Store listing complete',
        'Content rating done',
        'AAB uploaded',
        'Submitted for review'
    ]
    for item in google_items:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Apple App Store', 2)
    apple_items = [
        'Developer account created',
        'D-U-N-S number obtained',
        'Certificates created',
        'App Store Connect setup',
        'App uploaded via Xcode',
        'Submitted for review'
    ]
    for item in apple_items:
        doc.add_paragraph('☐ ' + item)
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_APP_STORE_GUIDE.docx')
    doc.save(output_path)
    print(f"App store guide saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_app_store_guide()
