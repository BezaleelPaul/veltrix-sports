"""
Convert sprint plan to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sprint_plan():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('2 Week Sprint Plan', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sprint info
    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.add_run('Duration: ').bold = True
    info.add_run('2 weeks (10 working days)\n')
    info.add_run('Goal: ').bold = True
    info.add_run('Working MVP with Login, Dashboard, Plans, Events\n')
    info.add_run('Team: ').bold = True
    info.add_run('2 Flutter + 1 Backend\n')
    info.add_run('Budget: ').bold = True
    info.add_run('₹30,000')
    
    doc.add_page_break()
    
    # DELIVERABLES
    doc.add_heading('DELIVERABLES', 1)
    
    doc.add_heading('Must Have (P0)', 2)
    must_have = [
        'User registration (email + phone)',
        'User login (email + password)',
        'OTP verification',
        'Dashboard with stats',
        'Training plans list',
        'Plan detail view',
        'Events list',
        'Event detail view',
        'Event registration',
        'Basic payment (Razorpay)'
    ]
    for item in must_have:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Skip (P2)', 2)
    skip_items = [
        'AI plan generation',
        'Device sync',
        'Ticketing marketplace',
        'Admin panel',
        'Analytics'
    ]
    for item in skip_items:
        doc.add_paragraph('✗ ' + item)
    
    doc.add_page_break()
    
    # WEEK 1
    doc.add_heading('WEEK 1: FOUNDATION', 1)
    
    # Day 1
    doc.add_heading('Day 1 (Monday) - Backend Setup', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day1_morning = [
        'Create Node.js project',
        'Setup Express + middleware',
        'Setup PostgreSQL connection',
        'Create database schema',
        'Run migrations'
    ]
    for item in day1_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day1_afternoon = [
        'Create User model',
        'Create Auth routes',
        'Implement register API',
        'Implement login API',
        'Test APIs with Postman'
    ]
    for item in day1_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 1 Deliverables: ').bold = True
    p.add_run('Backend project running, Database connected, Register/Login APIs working')
    
    # Day 2
    doc.add_heading('Day 2 (Tuesday) - Auth APIs + Flutter Setup', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day2_morning = [
        'Implement OTP generation',
        'Implement OTP verification',
        'Implement JWT token refresh',
        'Create Flutter project',
        'Setup project structure'
    ]
    for item in day2_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day2_afternoon = [
        'Add dependencies (dio, bloc, go_router)',
        'Create API client',
        'Create Auth repository',
        'Create Auth BLoC',
        'Test API connection'
    ]
    for item in day2_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 2 Deliverables: ').bold = True
    p.add_run('All auth APIs complete, Flutter project initialized, API connection working')
    
    # Day 3
    doc.add_heading('Day 3 (Wednesday) - Auth Screens', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day3_morning = [
        'Create Login screen UI',
        'Create Register screen UI',
        'Create OTP verification screen',
        'Add form validation',
        'Connect to Auth BLoC'
    ]
    for item in day3_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day3_afternoon = [
        'Implement login flow',
        'Implement register flow',
        'Implement OTP flow',
        'Store JWT token',
        'Test complete auth flow'
    ]
    for item in day3_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 3 Deliverables: ').bold = True
    p.add_run('Login screen working, Register screen working, OTP verification working')
    
    # Day 4
    doc.add_heading('Day 4 (Thursday) - Dashboard API + Screen', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day4_morning = [
        'Create Dashboard API',
        'Create User profile API',
        'Create Plans list API',
        'Create Events list API',
        'Test all APIs'
    ]
    for item in day4_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day4_afternoon = [
        'Create Dashboard screen UI',
        'Create stats cards',
        'Create quick actions',
        'Connect to APIs',
        'Test dashboard'
    ]
    for item in day4_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 4 Deliverables: ').bold = True
    p.add_run('Dashboard APIs complete, Dashboard screen working, User stats displayed')
    
    # Day 5
    doc.add_heading('Day 5 (Friday) - Training Plans', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day5_morning = [
        'Create Plans list screen',
        'Create Plan card widget',
        'Create Plan detail screen',
        'Create Session list widget',
        'Connect to Plans API'
    ]
    for item in day5_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day5_afternoon = [
        'Create Session detail screen',
        'Implement pull-to-refresh',
        'Add loading states',
        'Add error handling',
        'Test complete flow'
    ]
    for item in day5_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 5 Deliverables: ').bold = True
    p.add_run('Plans list working, Plan detail working, Session detail working, Week 1 complete')
    
    doc.add_page_break()
    
    # WEEK 2
    doc.add_heading('WEEK 2: FEATURES + POLISH', 1)
    
    # Day 6
    doc.add_heading('Day 6 (Monday) - Events', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day6_morning = [
        'Create Events list screen',
        'Create Event card widget',
        'Create Event detail screen',
        'Create filter/search UI',
        'Connect to Events API'
    ]
    for item in day6_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day6_afternoon = [
        'Implement event filters',
        'Implement search',
        'Add event images',
        'Create registration form',
        'Test events flow'
    ]
    for item in day6_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 6 Deliverables: ').bold = True
    p.add_run('Events list working, Event detail working, Filters working')
    
    # Day 7
    doc.add_heading('Day 7 (Tuesday) - Event Registration', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day7_morning = [
        'Create Registration API',
        'Create Categories API',
        'Create My Registrations API',
        'Implement registration logic',
        'Test APIs'
    ]
    for item in day7_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day7_afternoon = [
        'Complete registration form',
        'Add category selection',
        'Add form validation',
        'Create My Registrations screen',
        'Test registration flow'
    ]
    for item in day7_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 7 Deliverables: ').bold = True
    p.add_run('Registration API working, Registration form working, User can register for events')
    
    # Day 8
    doc.add_heading('Day 8 (Wednesday) - Payments', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day8_morning = [
        'Create Razorpay account',
        'Setup Razorpay keys',
        'Create Payment API',
        'Implement order creation',
        'Implement payment verification'
    ]
    for item in day8_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day8_afternoon = [
        'Add razorpay_flutter package',
        'Create Payment service',
        'Implement checkout flow',
        'Add payment success/failure handling',
        'Test payment flow'
    ]
    for item in day8_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 8 Deliverables: ').bold = True
    p.add_run('Razorpay integrated, Payment flow working, Payment verification working')
    
    # Day 9
    doc.add_heading('Day 9 (Thursday) - Testing + Bug Fixes', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day9_morning = [
        'Test complete user flow',
        'Test edge cases',
        'Fix critical bugs',
        'Fix UI issues',
        'Performance testing'
    ]
    for item in day9_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day9_afternoon = [
        'Fix remaining bugs',
        'Add loading indicators',
        'Add error messages',
        'Add empty states',
        'Final testing'
    ]
    for item in day9_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 9 Deliverables: ').bold = True
    p.add_run('All critical bugs fixed, UI polished, Ready for deployment')
    
    # Day 10
    doc.add_heading('Day 10 (Friday) - Deploy + Handover', 2)
    
    doc.add_heading('Morning (9 AM - 1 PM)', 3)
    day10_morning = [
        'Setup AWS EC2 instance',
        'Deploy backend API',
        'Setup PostgreSQL on RDS',
        'Configure environment vars',
        'Test deployed API'
    ]
    for item in day10_morning:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Afternoon (2 PM - 6 PM)', 3)
    day10_afternoon = [
        'Build Flutter APK/IPA',
        'Test on real devices',
        'Create release notes',
        'Handover documentation',
        'Sprint review meeting'
    ]
    for item in day10_afternoon:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Day 10 Deliverables: ').bold = True
    p.add_run('Backend deployed, App built and tested, Documentation complete, Handover done')
    
    doc.add_page_break()
    
    # API ENDPOINTS
    doc.add_heading('API ENDPOINTS', 1)
    
    doc.add_heading('Auth APIs', 2)
    auth_apis = [
        'POST /api/auth/register',
        'POST /api/auth/login',
        'POST /api/auth/otp/send',
        'POST /api/auth/otp/verify',
        'POST /api/auth/refresh'
    ]
    for api in auth_apis:
        doc.add_paragraph(api, style='List Bullet')
    
    doc.add_heading('User APIs', 2)
    user_apis = [
        'GET /api/user/profile',
        'PUT /api/user/profile'
    ]
    for api in user_apis:
        doc.add_paragraph(api, style='List Bullet')
    
    doc.add_heading('Training APIs', 2)
    training_apis = [
        'GET /api/plans',
        'GET /api/plans/:id',
        'GET /api/plans/:id/sessions',
        'GET /api/sessions/:id'
    ]
    for api in training_apis:
        doc.add_paragraph(api, style='List Bullet')
    
    doc.add_heading('Event APIs', 2)
    event_apis = [
        'GET /api/events',
        'GET /api/events/:id',
        'POST /api/events/:id/register',
        'GET /api/events/registrations'
    ]
    for api in event_apis:
        doc.add_paragraph(api, style='List Bullet')
    
    doc.add_heading('Payment APIs', 2)
    payment_apis = [
        'POST /api/payments/create',
        'POST /api/payments/verify'
    ]
    for api in payment_apis:
        doc.add_paragraph(api, style='List Bullet')
    
    doc.add_page_break()
    
    # FLUTTER SCREENS
    doc.add_heading('FLUTTER SCREENS', 1)
    
    doc.add_heading('Auth Flow', 2)
    auth_screens = [
        '1. Splash Screen',
        '2. Login Screen',
        '3. Register Screen',
        '4. OTP Screen'
    ]
    for screen in auth_screens:
        doc.add_paragraph(screen, style='List Number')
    
    doc.add_heading('Main Flow', 2)
    main_screens = [
        '5. Dashboard Screen',
        '6. Training Plans Screen',
        '7. Plan Detail Screen',
        '8. Session Detail Screen',
        '9. Events Screen',
        '10. Event Detail Screen',
        '11. Registration Screen',
        '12. My Registrations Screen',
        '13. Profile Screen'
    ]
    for screen in main_screens:
        doc.add_paragraph(screen, style='List Number')
    
    doc.add_page_break()
    
    # TECHNICAL STACK
    doc.add_heading('TECHNICAL STACK', 1)
    
    doc.add_heading('Backend', 2)
    backend_stack = [
        'Node.js + Express',
        'PostgreSQL',
        'JWT Authentication',
        'Razorpay SDK'
    ]
    for item in backend_stack:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Frontend', 2)
    frontend_stack = [
        'Flutter 3.41.9',
        'BLoC (State Management)',
        'GoRouter (Navigation)',
        'Dio (HTTP Client)'
    ]
    for item in frontend_stack:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Infrastructure', 2)
    infra_stack = [
        'AWS EC2 (Hosting)',
        'AWS RDS (Database)',
        'Razorpay (Payments)'
    ]
    for item in infra_stack:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # SUCCESS CRITERIA
    doc.add_heading('SUCCESS CRITERIA', 1)
    
    success_items = [
        'User can register and login',
        'User can view dashboard',
        'User can browse training plans',
        'User can browse events',
        'User can register for events',
        'User can make payment',
        'App deployed and working'
    ]
    for item in success_items:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Signatures
    doc.add_heading('APPROVAL', 1)
    
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    sig_data = [
        ('Sprint Start:', '_______________'),
        ('Sprint End:', '_______________'),
        ('Project Manager:', '_______________'),
        ('Tech Lead:', '_______________')
    ]
    for i, (key, value) in enumerate(sig_data):
        sig_table.cell(i, 0).text = key
        sig_table.cell(i, 1).text = value
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_SPRINT_PLAN.docx')
    doc.save(output_path)
    print(f"Sprint plan saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_sprint_plan()
