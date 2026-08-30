"""
Convert final showcase document to Word
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_final():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Project Proposal & MVP Specification', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Document Info
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Document Version', '1.0'),
        ('Date', 'August 29, 2026'),
        ('Classification', 'Confidential — For Internal Review'),
        ('Prepared For', 'Project Stakeholder'),
        ('Prepared By', 'Veltrix Sports Development Team'),
        ('Status', 'Pending Approval')
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    
    doc.add_paragraph(
        'Veltrix Sports is a mobile-first sports platform that connects athletes, coaches, '
        'event organizers, and spectators through three integrated modules: Training Plans, '
        'Events, and Ticketing.'
    )
    
    doc.add_paragraph(
        'The Minimum Viable Product (MVP) delivers a functional mobile application on Android '
        'and iOS, enabling users to browse and purchase training plans from certified coaches, '
        'discover and register for sports events, and buy tickets for tournaments and matches '
        'with integrated payment processing.'
    )
    
    doc.add_heading('MVP Highlights', 2)
    highlights_table = doc.add_table(rows=8, cols=2)
    highlights_table.style = 'Table Grid'
    highlights_data = [
        ('Item', 'Details'),
        ('Target Users', 'Athletes, Coaches, Event Organizers, Spectators'),
        ('Core Modules', 'Training Plans, Events, Ticketing'),
        ('Target Platforms', 'Android (Primary), iOS (Secondary)'),
        ('MVP Timeline', '2 weeks (10 business days)'),
        ('Estimated Budget', '₹91,800 (Development + Infrastructure + Services + App Stores)'),
        ('Primary Backend', '[REQUIRES CONFIRMATION]'),
        ('Payment Gateway', 'Razorpay')
    ]
    for i, (key, value) in enumerate(highlights_data):
        highlights_table.cell(i, 0).text = key
        highlights_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. Project Overview
    doc.add_heading('2. Project Overview', 1)
    
    doc.add_heading('2.1 Problem Statement', 2)
    doc.add_paragraph(
        'Athletes and sports enthusiasts lack a unified platform to discover coaching services, '
        'find local sports events, and purchase event tickets. Current solutions are fragmented '
        'across multiple apps and platforms.'
    )
    
    doc.add_heading('2.2 Solution', 2)
    doc.add_paragraph(
        'Veltrix Sports provides a single mobile application that integrates Training (access to '
        'certified coaches and structured training programs), Events (discovery and registration '
        'for sports events), and Ticketing (seamless ticket purchasing with digital delivery).'
    )
    
    doc.add_heading('2.3 User Groups', 2)
    user_table = doc.add_table(rows=5, cols=2)
    user_table.style = 'Table Grid'
    user_data = [
        ('User Group', 'Primary Actions'),
        ('Athletes', 'Browse training plans, track progress, register for events, purchase tickets'),
        ('Coaches', 'Create training plans, manage sessions, organize events'),
        ('Event Organizers', 'Create events, manage registrations, sell tickets'),
        ('Spectators', 'Browse events, purchase tickets, receive digital tickets')
    ]
    for i, (key, value) in enumerate(user_data):
        user_table.cell(i, 0).text = key
        user_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 3. Product Scope
    doc.add_heading('3. Product Scope', 1)
    
    doc.add_heading('3.1 MVP Scope Statement', 2)
    scope_table = doc.add_table(rows=9, cols=2)
    scope_table.style = 'Table Grid'
    scope_data = [
        ('Module', 'Included in MVP'),
        ('User Authentication', 'Yes'),
        ('Training Plans (Browse, Purchase, Track)', 'Yes'),
        ('Event Discovery & Registration', 'Yes'),
        ('Ticket Purchasing & Digital Delivery', 'Yes'),
        ('User Profile Management', 'Yes'),
        ('Payment Processing (Razorpay)', 'Yes'),
        ('Push Notifications', 'Yes'),
        ('QR Code Generation & Scanning', 'Yes')
    ]
    for i, (key, value) in enumerate(scope_data):
        scope_table.cell(i, 0).text = key
        scope_table.cell(i, 1).text = value
    
    doc.add_heading('3.2 Platform Scope', 2)
    platform_table = doc.add_table(rows=4, cols=2)
    platform_table.style = 'Table Grid'
    platform_data = [
        ('Platform', 'MVP Delivery'),
        ('Android', 'Primary target'),
        ('iOS', 'Secondary target — [REQUIRES CONFIRMATION]'),
        ('Web', '[REQUIRES CONFIRMATION — not included unless specified]')
    ]
    for i, (key, value) in enumerate(platform_data):
        platform_table.cell(i, 0).text = key
        platform_table.cell(i, 1).text = value
    
    doc.add_paragraph(
        'Note: The 2-week MVP timeline assumes Android as the primary platform. iOS delivery '
        'will follow Android completion. Web platform scope requires stakeholder confirmation.'
    )
    
    doc.add_page_break()
    
    # 4. Core Modules
    doc.add_heading('4. Core Modules', 1)
    
    doc.add_heading('4.1 Training Plans', 2)
    training_features = [
        ('Coach Profiles', 'View coach information, ratings, specializations'),
        ('Training Plan Catalog', 'Browse plans by sport, level, price'),
        ('Plan Details', 'View schedule, curriculum, reviews'),
        ('Purchase & Enrollment', 'Buy plans via integrated payment'),
        ('Video Sessions', 'Access training videos within sessions'),
        ('Progress Tracking', 'Track completed sessions, streaks, hours')
    ]
    training_table = doc.add_table(rows=len(training_features)+1, cols=2)
    training_table.style = 'Table Grid'
    training_table.cell(0, 0).text = 'Feature'
    training_table.cell(0, 1).text = 'Description'
    for i, (feature, desc) in enumerate(training_features):
        training_table.cell(i+1, 0).text = feature
        training_table.cell(i+1, 1).text = desc
    
    doc.add_heading('4.2 Events', 2)
    event_features = [
        ('Event Discovery', 'Browse events by sport, location, date'),
        ('Event Details', 'View venue, organizer, rules, prizes'),
        ('Registration', 'Register for events with participant details'),
        ('QR Check-in', 'Digital check-in via QR code'),
        ('Calendar Integration', 'Add events to device calendar')
    ]
    event_table = doc.add_table(rows=len(event_features)+1, cols=2)
    event_table.style = 'Table Grid'
    event_table.cell(0, 0).text = 'Feature'
    event_table.cell(0, 1).text = 'Description'
    for i, (feature, desc) in enumerate(event_features):
        event_table.cell(i+1, 0).text = feature
        event_table.cell(i+1, 1).text = desc
    
    doc.add_heading('4.3 Ticketing', 2)
    ticket_features = [
        ('Ticket Catalog', 'Browse available tickets for events'),
        ('Seat Selection', 'Interactive seat map for venue selection'),
        ('Ticket Purchase', 'Buy tickets with multiple payment options'),
        ('Digital Tickets', 'QR code-based tickets with download option'),
        ('Ticket Transfer', 'Transfer tickets to other users')
    ]
    ticket_table = doc.add_table(rows=len(ticket_features)+1, cols=2)
    ticket_table.style = 'Table Grid'
    ticket_table.cell(0, 0).text = 'Feature'
    ticket_table.cell(0, 1).text = 'Description'
    for i, (feature, desc) in enumerate(ticket_features):
        ticket_table.cell(i+1, 0).text = feature
        ticket_table.cell(i+1, 1).text = desc
    
    doc.add_page_break()
    
    # 5. App Screens & User Flows
    doc.add_heading('5. App Screens & User Flows', 1)
    
    doc.add_heading('5.1 Screen Inventory', 2)
    doc.add_paragraph('The MVP includes 28 screens across the following modules:')
    
    screen_table = doc.add_table(rows=10, cols=3)
    screen_table.style = 'Table Grid'
    screen_data = [
        ('Module', 'Screen Count', 'Screens'),
        ('Onboarding', '4', 'Splash, Welcome 1, Welcome 2, Welcome 3'),
        ('Authentication', '4', 'Login, Sign Up, Forgot Password, OTP Verification'),
        ('Home', '1', 'Dashboard with quick actions and featured content'),
        ('Training', '4', 'Plans List, Plan Details, Session Player, Progress Dashboard'),
        ('Events', '4', 'Events List, Event Details, Registration, Check-in'),
        ('Tickets', '4', 'Tickets List, Ticket Details, Seat Selection, Booking Confirmation'),
        ('Cart & Payment', '4', 'Cart, Checkout, Payment Success, Payment Failed'),
        ('Profile', '4', 'Profile, Edit Profile, My Bookings, My Tickets'),
        ('Coaches', '3', 'Coach List, Coach Profile, Book Session')
    ]
    for i, row_data in enumerate(screen_data):
        for j, cell_data in enumerate(row_data):
            screen_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 6. Functional Requirements
    doc.add_heading('6. Functional Requirements', 1)
    
    doc.add_heading('6.1 Authentication & Authorization', 2)
    auth_req_table = doc.add_table(rows=6, cols=3)
    auth_req_table.style = 'Table Grid'
    auth_req_data = [
        ('ID', 'Requirement', 'Priority'),
        ('AUTH-01', 'Users can register with email and password', 'High'),
        ('AUTH-02', 'Users can sign in with Google account', 'High'),
        ('AUTH-03', 'Users can sign in with Apple ID', 'High'),
        ('AUTH-04', 'Users can reset password via email', 'High'),
        ('AUTH-05', 'Users can verify phone via OTP', 'Medium')
    ]
    for i, row_data in enumerate(auth_req_data):
        for j, cell_data in enumerate(row_data):
            auth_req_table.cell(i, j).text = cell_data
    
    doc.add_heading('6.2 Training Plans', 2)
    trn_req_table = doc.add_table(rows=6, cols=3)
    trn_req_table.style = 'Table Grid'
    trn_req_data = [
        ('ID', 'Requirement', 'Priority'),
        ('TRN-01', 'Users can browse training plans with search and filter', 'High'),
        ('TRN-02', 'Users can view plan details including schedule and reviews', 'High'),
        ('TRN-03', 'Users can purchase plans via integrated payment', 'High'),
        ('TRN-04', 'Users can access video sessions within purchased plans', 'High'),
        ('TRN-05', 'Users can track progress including sessions and streaks', 'High')
    ]
    for i, row_data in enumerate(trn_req_data):
        for j, cell_data in enumerate(row_data):
            trn_req_table.cell(i, j).text = cell_data
    
    doc.add_heading('6.3 Events', 2)
    evt_req_table = doc.add_table(rows=6, cols=3)
    evt_req_table.style = 'Table Grid'
    evt_req_data = [
        ('ID', 'Requirement', 'Priority'),
        ('EVT-01', 'Users can browse events with search and filter', 'High'),
        ('EVT-02', 'Users can view event details including venue and rules', 'High'),
        ('EVT-03', 'Users can register for events with participant details', 'High'),
        ('EVT-04', 'Users can check in via QR code', 'Medium'),
        ('EVT-05', 'Users can add events to device calendar', 'Medium')
    ]
    for i, row_data in enumerate(evt_req_data):
        for j, cell_data in enumerate(row_data):
            evt_req_table.cell(i, j).text = cell_data
    
    doc.add_heading('6.4 Ticketing', 2)
    tkt_req_table = doc.add_table(rows=8, cols=3)
    tkt_req_table.style = 'Table Grid'
    tkt_req_data = [
        ('ID', 'Requirement', 'Priority'),
        ('TKT-01', 'Users can browse available tickets', 'High'),
        ('TKT-02', 'Users can view ticket details and pricing', 'High'),
        ('TKT-03', 'Users can select seats via interactive map', 'High'),
        ('TKT-04', 'Users can purchase tickets via integrated payment', 'High'),
        ('TKT-05', 'Users receive QR code tickets', 'High'),
        ('TKT-06', 'Users can download tickets', 'Medium'),
        ('TKT-07', 'Users can transfer tickets to other users', 'Medium')
    ]
    for i, row_data in enumerate(tkt_req_data):
        for j, cell_data in enumerate(row_data):
            tkt_req_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 7. Technical Architecture
    doc.add_heading('7. Technical Architecture', 1)
    
    doc.add_heading('7.1 Architecture Overview', 2)
    doc.add_paragraph(
        'The application follows a layered architecture pattern with clear separation of concerns '
        'between Presentation (Flutter UI), Business Logic (BLoC), Data Layer (Repositories, '
        'Data Sources), Network Layer (Dio), and Backend Services.'
    )
    
    doc.add_heading('7.2 Backend Architecture', 2)
    doc.add_paragraph(
        '[BACKEND ARCHITECTURE REQUIRES CONFIRMATION]'
    )
    doc.add_paragraph(
        'The source documentation references both Firebase and AWS without clearly defining '
        'service boundaries. The following options require stakeholder confirmation:'
    )
    
    backend_options = [
        'Option A: Firebase-Primary (Firebase Auth, Cloud Firestore, Firebase Storage)',
        'Option B: AWS-Primary (AWS Cognito, AWS RDS PostgreSQL, AWS S3)',
        'Option C: Hybrid (Firebase Auth, PostgreSQL on AWS, AWS S3)'
    ]
    for option in backend_options:
        doc.add_paragraph(option, style='List Bullet')
    
    doc.add_paragraph('Action Required: Confirm backend architecture before development begins.')
    
    doc.add_page_break()
    
    # 8. Technology Stack
    doc.add_heading('8. Technology Stack', 1)
    
    tech_table = doc.add_table(rows=10, cols=3)
    tech_table.style = 'Table Grid'
    tech_data = [
        ('Layer', 'Technology', 'Purpose'),
        ('Frontend Framework', 'Flutter 3.41.9', 'Cross-platform mobile development'),
        ('Programming Language', 'Dart 3.11.5', 'Application logic'),
        ('State Management', 'BLoC', 'Application state'),
        ('Navigation', 'GoRouter', 'Screen navigation'),
        ('HTTP Client', 'Dio', 'API communication'),
        ('Local Storage', 'Hive', 'Local data persistence'),
        ('Backend', '[REQUIRES CONFIRMATION]', 'Backend services'),
        ('Database', '[REQUIRES CONFIRMATION]', 'Data persistence'),
        ('Payment Gateway', 'Razorpay', 'Payment processing')
    ]
    for i, row_data in enumerate(tech_data):
        for j, cell_data in enumerate(row_data):
            tech_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 9. Project Assumptions & Dependencies
    doc.add_heading('9. Project Assumptions & Dependencies', 1)
    
    doc.add_heading('9.1 Confirmed Assumptions', 2)
    confirmed_table = doc.add_table(rows=6, cols=3)
    confirmed_table.style = 'Table Grid'
    confirmed_data = [
        ('ID', 'Assumption', 'Impact if Invalid'),
        ('ASM-01', 'Flutter and Dart SDK are installed', 'Development cannot begin'),
        ('ASM-02', 'Razorpay merchant account is available', 'Payment integration blocked'),
        ('ASM-03', 'Development team has Flutter expertise', 'Timeline impacted'),
        ('ASM-04', 'Basic UI/UX designs are available', 'Development delayed'),
        ('ASM-05', '2-week timeline assumes dedicated team', 'Timeline extended')
    ]
    for i, row_data in enumerate(confirmed_data):
        for j, cell_data in enumerate(row_data):
            confirmed_table.cell(i, j).text = cell_data
    
    doc.add_heading('9.2 Assumptions Requiring Confirmation', 2)
    unconfirmed_table = doc.add_table(rows=7, cols=3)
    unconfirmed_table.style = 'Table Grid'
    unconfirmed_data = [
        ('ID', 'Assumption', 'Status'),
        ('ASM-06', 'Backend infrastructure details are confirmed', 'REQUIRES CONFIRMATION'),
        ('ASM-07', 'API documentation is available', 'REQUIRES CONFIRMATION'),
        ('ASM-08', 'Google Play Developer account is active', 'REQUIRES CONFIRMATION'),
        ('ASM-09', 'Apple Developer account is active', 'REQUIRES CONFIRMATION'),
        ('ASM-10', 'iOS delivery is within MVP scope', 'REQUIRES CONFIRMATION'),
        ('ASM-11', 'Web delivery is within MVP scope', 'REQUIRES CONFIRMATION')
    ]
    for i, row_data in enumerate(unconfirmed_data):
        for j, cell_data in enumerate(row_data):
            unconfirmed_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 10. Development Timeline
    doc.add_heading('10. Development Timeline', 1)
    
    doc.add_heading('10.1 Development Phase (2 Weeks / 10 Business Days)', 2)
    timeline_table = doc.add_table(rows=11, cols=3)
    timeline_table.style = 'Table Grid'
    timeline_data = [
        ('Day', 'Focus', 'Deliverables'),
        ('1', 'Backend Setup', 'Database schema, Auth APIs, Project scaffolding'),
        ('2', 'Flutter Setup', 'Project structure, Theme system, Navigation setup'),
        ('3', 'Authentication', 'Login, Signup, Forgot Password, OTP screens'),
        ('4', 'Home Dashboard', 'Home screen, Quick actions, Featured content'),
        ('5', 'Training Module', 'Plans list, Plan details, Session player'),
        ('6', 'Events Module', 'Events list, Event details, Registration'),
        ('7', 'Ticketing Module', 'Tickets list, Seat selection, Booking'),
        ('8', 'Payments', 'Razorpay integration, Cart, Checkout flow'),
        ('9', 'Testing & Polish', 'Bug fixes, Performance optimization, UI refinement'),
        ('10', 'Deployment', 'Build generation, Store upload, Documentation')
    ]
    for i, row_data in enumerate(timeline_data):
        for j, cell_data in enumerate(row_data):
            timeline_table.cell(i, j).text = cell_data
    
    doc.add_heading('10.2 App Store Review', 2)
    doc.add_paragraph(
        'Store review and publishing timelines are dependent on platform review processes '
        'and are outside the direct control of the development team.'
    )
    
    store_table = doc.add_table(rows=3, cols=2)
    store_table.style = 'Table Grid'
    store_data = [
        ('Store', 'Typical Review Time'),
        ('Google Play', '1-7 days'),
        ('Apple App Store', '1-2 days')
    ]
    for i, (key, value) in enumerate(store_data):
        store_table.cell(i, 0).text = key
        store_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 11. QA & Testing Strategy
    doc.add_heading('11. QA & Testing Strategy', 1)
    
    doc.add_heading('11.1 Testing Approach', 2)
    testing_table = doc.add_table(rows=11, cols=3)
    testing_table.style = 'Table Grid'
    testing_data = [
        ('Test Type', 'Scope', 'Priority'),
        ('Functional Testing', 'All user-facing features', 'High'),
        ('UI Testing', 'Screen layouts, responsiveness', 'High'),
        ('API Integration', 'Backend communication', 'High'),
        ('Authentication', 'Login, signup, password reset', 'High'),
        ('Payment Flow', 'Razorpay integration, success/failure', 'High'),
        ('QR Functionality', 'Ticket generation, scan validation', 'Medium'),
        ('Error Handling', 'Network errors, validation errors', 'High'),
        ('Regression', 'Re-testing after bug fixes', 'Medium'),
        ('Performance', 'Load times, memory usage', 'Medium'),
        ('Device Testing', 'Representative device configurations', 'Medium')
    ]
    for i, row_data in enumerate(testing_data):
        for j, cell_data in enumerate(row_data):
            testing_table.cell(i, j).text = cell_data
    
    doc.add_heading('11.2 Device Testing Strategy', 2)
    doc.add_paragraph(
        'Testing will cover representative supported devices, screen sizes, OS versions, '
        'and release configurations. Full device matrix testing is outside MVP scope.'
    )
    
    doc.add_heading('11.3 Defect Classification', 2)
    defect_table = doc.add_table(rows=5, cols=3)
    defect_table.style = 'Table Grid'
    defect_data = [
        ('Severity', 'Description', 'Resolution'),
        ('Critical', 'App crashes, data loss, payment failure', 'Must fix before release'),
        ('Major', 'Feature not working, UI broken', 'Should fix before release'),
        ('Minor', 'Cosmetic issues, minor UX problems', 'Can defer to post-release'),
        ('Trivial', 'Typos, alignment issues', 'Can defer to post-release')
    ]
    for i, row_data in enumerate(defect_data):
        for j, cell_data in enumerate(row_data):
            defect_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 12. Deployment & App Store Publishing
    doc.add_heading('12. Deployment & App Store Publishing', 1)
    
    doc.add_heading('12.1 Android (Google Play Store)', 2)
    android_table = doc.add_table(rows=7, cols=3)
    android_table.style = 'Table Grid'
    android_data = [
        ('Step', 'Description', 'Timeline'),
        ('1', 'Build AAB file', 'Day 10'),
        ('2', 'Upload to Google Play Console', 'Day 10'),
        ('3', 'Complete store listing', 'Day 10'),
        ('4', 'Submit for review', 'Day 10'),
        ('5', 'Google review process', '1-7 days'),
        ('6', 'App published', 'Upon approval')
    ]
    for i, row_data in enumerate(android_data):
        for j, cell_data in enumerate(row_data):
            android_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('Cost: ₹18,000 (one-time developer fee)')
    
    doc.add_heading('12.2 iOS (Apple App Store)', 2)
    doc.add_paragraph(
        'If iOS is included in MVP scope, additional development time may be required '
        'beyond the 2-week timeline.'
    )
    ios_table = doc.add_table(rows=7, cols=3)
    ios_table.style = 'Table Grid'
    ios_data = [
        ('Step', 'Description', 'Timeline'),
        ('1', 'Build IPA file via Xcode', 'Day 10'),
        ('2', 'Upload to App Store Connect', 'Day 10'),
        ('3', 'Complete store listing', 'Day 10'),
        ('4', 'Submit for review', 'Day 10'),
        ('5', 'Apple review process', '1-2 days'),
        ('6', 'App published', 'Upon approval')
    ]
    for i, row_data in enumerate(ios_data):
        for j, cell_data in enumerate(row_data):
            ios_table.cell(i, j).text = cell_data
    
    doc.add_paragraph('Cost: ₹7,500/year (developer fee)')
    
    doc.add_page_break()
    
    # 13. Team Requirements
    doc.add_heading('13. Team Requirements', 1)
    
    doc.add_heading('13.1 Team Composition', 2)
    team_table = doc.add_table(rows=5, cols=3)
    team_table.style = 'Table Grid'
    team_data = [
        ('Role', 'Count', 'Responsibilities'),
        ('Flutter Developer', '2', 'Frontend development, UI implementation'),
        ('Backend Developer', '1', 'API development, database, server logic'),
        ('UI/UX Designer', '1', 'Design deliverables (assumed available)'),
        ('QA Tester', '1', 'Testing and quality assurance')
    ]
    for i, row_data in enumerate(team_data):
        for j, cell_data in enumerate(row_data):
            team_table.cell(i, j).text = cell_data
    
    doc.add_heading('13.2 Responsibilities', 2)
    resp_table = doc.add_table(rows=4, cols=2)
    resp_table.style = 'Table Grid'
    resp_data = [
        ('Responsibility', 'Assigned To'),
        ('Flutter App Development', 'User + Team'),
        ('Project Planning & Coordination', 'User'),
        ('UI/UX Design', 'Design Team (assumed)')
    ]
    for i, (key, value) in enumerate(resp_data):
        resp_table.cell(i, 0).text = key
        resp_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 14. Cost Breakdown
    doc.add_heading('14. Cost Breakdown', 1)
    
    doc.add_heading('14.1 Development Costs', 2)
    dev_table = doc.add_table(rows=5, cols=2)
    dev_table.style = 'Table Grid'
    dev_data = [
        ('Item', 'Cost (₹)'),
        ('Flutter Development (2 developers × 10 days)', '40,000'),
        ('Backend Development (1 developer × 5 days)', '15,000'),
        ('UI/UX Design', '10,000'),
        ('Development Subtotal', '65,000')
    ]
    for i, (key, value) in enumerate(dev_data):
        dev_table.cell(i, 0).text = key
        dev_table.cell(i, 1).text = value
    
    doc.add_heading('14.2 Infrastructure Costs', 2)
    infra_table = doc.add_table(rows=3, cols=2)
    infra_table.style = 'Table Grid'
    infra_data = [
        ('Item', 'Cost (₹)'),
        ('AWS (Free Tier)', '0'),
        ('Domain Name', '800')
    ]
    for i, (key, value) in enumerate(infra_data):
        infra_table.cell(i, 0).text = key
        infra_table.cell(i, 1).text = value
    
    doc.add_heading('14.3 Third-Party Services', 2)
    services_table = doc.add_table(rows=4, cols=2)
    services_table.style = 'Table Grid'
    services_data = [
        ('Item', 'Cost (₹)'),
        ('Firebase', '0 (Free tier)'),
        ('Razorpay', '2% per transaction (variable)'),
        ('Twilio (OTP)', '500')
    ]
    for i, (key, value) in enumerate(services_data):
        services_table.cell(i, 0).text = key
        services_table.cell(i, 1).text = value
    
    doc.add_heading('14.4 App Store Fees', 2)
    store_fee_table = doc.add_table(rows=3, cols=2)
    store_fee_table.style = 'Table Grid'
    store_fee_data = [
        ('Item', 'Cost (₹)'),
        ('Google Play Developer', '18,000 (one-time)'),
        ('Apple Developer', '7,500/year')
    ]
    for i, (key, value) in enumerate(store_fee_data):
        store_fee_table.cell(i, 0).text = key
        store_fee_table.cell(i, 1).text = value
    
    doc.add_heading('14.5 Total Investment', 2)
    total_table = doc.add_table(rows=6, cols=2)
    total_table.style = 'Table Grid'
    total_data = [
        ('Category', 'Amount (₹)'),
        ('Development', '65,000'),
        ('Infrastructure', '800'),
        ('Services', '500'),
        ('App Stores', '25,500'),
        ('GRAND TOTAL', '91,800')
    ]
    for i, (key, value) in enumerate(total_data):
        total_table.cell(i, 0).text = key
        total_table.cell(i, 1).text = value
    
    doc.add_paragraph(
        'Note: Razorpay transaction fees are variable and not included in the fixed cost estimate.'
    )
    
    doc.add_page_break()
    
    # 15. Out of Scope for MVP
    doc.add_heading('15. Out of Scope for MVP', 1)
    
    doc.add_paragraph(
        'The following features are explicitly excluded from the MVP and will be treated as '
        'change requests if required:'
    )
    
    oos_table = doc.add_table(rows=8, cols=2)
    oos_table.style = 'Table Grid'
    oos_data = [
        ('Category', 'Excluded Features'),
        ('Training', 'Live video streaming, Real-time coaching, Social features'),
        ('Events', 'Live scoring, Spectator chat, Event streaming'),
        ('Ticketing', 'Dynamic pricing, Auction tickets, Group discounts'),
        ('Platform', 'Web application, Desktop application'),
        ('Integrations', 'Fitness device sync, Social media login beyond Google/Apple'),
        ('Features', 'In-app chat, Video calls, Advanced analytics, AI recommendations'),
        ('Admin', 'Admin dashboard, Content management system, Analytics dashboard')
    ]
    for i, (key, value) in enumerate(oos_data):
        oos_table.cell(i, 0).text = key
        oos_table.cell(i, 1).text = value
    
    doc.add_paragraph(
        'Scope Control: Features not explicitly listed in the approved MVP scope will be '
        'treated as change requests and may affect timeline and cost.'
    )
    
    doc.add_page_break()
    
    # 16. MVP Acceptance Criteria
    doc.add_heading('16. MVP Acceptance Criteria', 1)
    
    doc.add_paragraph(
        'The MVP will be considered complete and ready for deployment when:'
    )
    
    ac_table = doc.add_table(rows=13, cols=2)
    ac_table.style = 'Table Grid'
    ac_data = [
        ('ID', 'Criterion'),
        ('AC-01', 'All 28 screens are implemented per specification'),
        ('AC-02', 'User authentication works (email, Google, Apple)'),
        ('AC-03', 'Training plan browsing, purchase, and playback function'),
        ('AC-04', 'Event browsing and registration function'),
        ('AC-05', 'Ticket browsing, seat selection, and purchase function'),
        ('AC-06', 'Razorpay payment integration functions correctly'),
        ('AC-07', 'QR code generation and display function'),
        ('AC-08', 'User profile management functions'),
        ('AC-09', 'All critical and major defects are resolved'),
        ('AC-10', 'Release build passes QA validation'),
        ('AC-11', 'App store submission is complete'),
        ('AC-12', 'Stakeholder review and approval is obtained')
    ]
    for i, (key, value) in enumerate(ac_data):
        ac_table.cell(i, 0).text = key
        ac_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 17. Risks & Mitigation
    doc.add_heading('17. Risks & Mitigation', 1)
    
    risk_table = doc.add_table(rows=8, cols=4)
    risk_table.style = 'Table Grid'
    risk_data = [
        ('Risk', 'Probability', 'Impact', 'Mitigation'),
        ('Scope creep', 'High', 'High', 'Strict adherence to approved scope'),
        ('Backend delays', 'Medium', 'High', 'Early API provisioning; mock services'),
        ('Payment integration issues', 'Medium', 'High', 'Early Razorpay sandbox testing'),
        ('App store rejection', 'Low', 'Medium', 'Follow platform guidelines'),
        ('Timeline overrun', 'Medium', 'High', 'Daily standups; scope prioritization'),
        ('Third-party service issues', 'Low', 'Medium', 'Fallback options; error handling'),
        ('Design delays', 'Medium', 'High', 'Early design delivery; component-based approach')
    ]
    for i, row_data in enumerate(risk_data):
        for j, cell_data in enumerate(row_data):
            risk_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 18. Immediate Next Steps
    doc.add_heading('18. Immediate Next Steps', 1)
    
    next_table = doc.add_table(rows=9, cols=4)
    next_table.style = 'Table Grid'
    next_data = [
        ('#', 'Action', 'Owner', 'Dependencies'),
        ('1', 'Review and approve this document', 'Stakeholder', 'None'),
        ('2', 'Confirm backend architecture', 'Stakeholder + Tech Lead', 'Section 7'),
        ('3', 'Confirm platform scope', 'Stakeholder', 'Section 3.2'),
        ('4', 'Provide UI/UX design deliverables', 'Design Team', 'Day 1'),
        ('5', 'Provision backend infrastructure', 'Backend Team', 'Day 1'),
        ('6', 'Configure third-party services', 'Backend Team', 'Day 1'),
        ('7', 'Verify app-store developer accounts', 'Business Team', 'Day 10'),
        ('8', 'Begin development', 'Development Team', 'Upon approval')
    ]
    for i, row_data in enumerate(next_data):
        for j, cell_data in enumerate(row_data):
            next_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 19. Conclusion
    doc.add_heading('19. Conclusion', 1)
    
    conclusion_table = doc.add_table(rows=8, cols=2)
    conclusion_table.style = 'Table Grid'
    conclusion_data = [
        ('Item', 'Status'),
        ('Documentation', 'Complete'),
        ('UI/UX Specification', 'Complete'),
        ('Cost Estimation', 'Complete'),
        ('Timeline', 'Defined (2 weeks)'),
        ('Team', 'Defined'),
        ('Backend Architecture', 'REQUIRES CONFIRMATION'),
        ('Platform Scope', 'REQUIRES CONFIRMATION')
    ]
    for i, (key, value) in enumerate(conclusion_data):
        conclusion_table.cell(i, 0).text = key
        conclusion_table.cell(i, 1).text = value
    
    doc.add_paragraph(
        'This document provides a comprehensive overview of the Veltrix Sports MVP. '
        'The project is ready to proceed upon stakeholder approval and resolution of '
        'items marked as requiring confirmation.'
    )
    
    doc.add_page_break()
    
    # Appendix: Change Summary
    doc.add_heading('Appendix: Change Summary', 1)
    
    doc.add_heading('Critical Changes Made', 2)
    critical_changes = [
        'Budget consistency: Aligned all budget figures to ₹91,800 total',
        'Backend architecture: Marked as requiring confirmation',
        'Platform scope: Clarified that iOS and Web require confirmation',
        'Removed unsupported claims: Removed "95% Android in India" statistic',
        'Added acceptance criteria: Defined clear completion criteria',
        'Added scope boundaries: Documented out-of-scope items',
        'Added assumptions: Documented confirmed and unconfirmed assumptions',
        'Improved QA section: Added comprehensive testing strategy',
        'Professional language: Removed marketing-style claims'
    ]
    for change in critical_changes:
        doc.add_paragraph(change, style='List Bullet')
    
    doc.add_heading('Items Requiring Stakeholder Confirmation', 2)
    confirmation_items = [
        'Backend Architecture: Firebase, AWS, or Hybrid? (Section 7.2)',
        'Platform Scope: Is iOS included in 2-week MVP? Is Web included? (Section 3.2)',
        'App Store Accounts: Are developer accounts active? (Section 9.2)',
        'UI/UX Designs: Are designs available or will they be created? (Section 9.3)',
        'API Documentation: Is backend API documentation available? (Section 9.3)',
        'Third-Party Credentials: Are Razorpay and OTP service credentials configured? (Section 9.3)'
    ]
    for item in confirmation_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Risks Identified', 2)
    risks = [
        'Scope creep — High probability, high impact',
        'Backend delays — Medium probability, high impact',
        'Timeline overrun — Medium probability, high impact',
        'App store rejection — Low probability, medium impact'
    ]
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_FINAL.docx')
    doc.save(output_path)
    print(f"Final document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_final()
