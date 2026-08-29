"""
Convert decision guide to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_decision_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('App Store Decision & Publishing Guide', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. DO YOU NEED BOTH?
    doc.add_heading('1. DO YOU NEED BOTH STORES?', 1)
    
    doc.add_heading('Quick Answer', 2)
    answer_table = doc.add_table(rows=4, cols=2)
    answer_table.style = 'Table Grid'
    answer_data = [
        ('Platform', 'Store Required'),
        ('Android Users', 'Google Play Store'),
        ('iOS Users', 'Apple App Store'),
        ('Web Users', 'None (direct URL)')
    ]
    for i, (key, value) in enumerate(answer_data):
        answer_table.cell(i, 0).text = key
        answer_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. YOUR OPTIONS
    doc.add_heading('2. YOUR OPTIONS', 1)
    
    doc.add_heading('Option A: Android Only (RECOMMENDED FOR MVP)', 2)
    option_a = doc.add_table(rows=4, cols=2)
    option_a.style = 'Table Grid'
    option_a_data = [
        ('Item', 'Details'),
        ('Cost', '₹18,000 (one-time)'),
        ('Time', '3-10 days'),
        ('Best For', 'MVP, Startups')
    ]
    for i, (key, value) in enumerate(option_a_data):
        option_a.cell(i, 0).text = key
        option_a.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Option B: iOS Only', 2)
    option_b = doc.add_table(rows=4, cols=2)
    option_b.style = 'Table Grid'
    option_b_data = [
        ('Item', 'Details'),
        ('Cost', '₹7,500/year'),
        ('Time', '10-20 days'),
        ('Best For', 'Premium market')
    ]
    for i, (key, value) in enumerate(option_b_data):
        option_b.cell(i, 0).text = key
        option_b.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Option C: Both Stores', 2)
    option_c = doc.add_table(rows=4, cols=2)
    option_c.style = 'Table Grid'
    option_c_data = [
        ('Item', 'Details'),
        ('Cost', '₹25,500 first year'),
        ('Time', '10-20 days'),
        ('Best For', 'Full launch')
    ]
    for i, (key, value) in enumerate(option_c_data):
        option_c.cell(i, 0).text = key
        option_c.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 3. RECOMMENDATION
    doc.add_heading('3. RECOMMENDATION', 1)
    
    doc.add_heading('For 2-Week MVP: Start with Android Only', 2)
    why_items = [
        'Faster: 3-10 days vs 10-20 days',
        'Cheaper: ₹18,000 vs ₹25,500',
        'More Users: 95% Android in India',
        'Easier: Simpler process'
    ]
    for item in why_items:
        doc.add_paragraph('✓ ' + item)
    
    doc.add_heading('Add iOS Later When:', 2)
    later_items = [
        'Product-market fit achieved',
        'Users request iOS app',
        'Revenue supports cost'
    ]
    for item in later_items:
        doc.add_paragraph('• ' + item)
    
    doc.add_page_break()
    
    # 4. GOOGLE PLAY - STEP BY STEP
    doc.add_heading('4. GOOGLE PLAY STORE - STEP BY STEP', 1)
    
    steps = [
        ('Step 1: Create Developer Account (Day 1)', [
            'Go to https://play.google.com/console',
            'Sign in with Gmail',
            'Click "Create Developer Account"',
            'Pay ₹18,000 (one-time)',
            'Fill developer name: "Veltrix Sports"',
            'Fill email and phone',
            'Verify email',
            'Wait 24-48 hours for approval'
        ]),
        ('Step 2: Create App (Day 2)', [
            'Go to Google Play Console',
            'Click "Create App"',
            'App name: "Veltrix Sports"',
            'Select: Free, App',
            'Accept declarations',
            'Click "Create App"'
        ]),
        ('Step 3: Complete Store Listing (Day 2-3)', [
            'Go to "Store presence" → "Main store listing"',
            'App name (max 30 chars)',
            'Short description (max 80 chars)',
            'Full description (max 4000 chars)',
            'Upload app icon (512x512 PNG)',
            'Upload feature graphic (1024x500 PNG)',
            'Upload screenshots (min 2, max 8)'
        ]),
        ('Step 4: Content Rating (Day 3)', [
            'Go to "Store presence" → "Content rating"',
            'Fill questionnaire',
            'Get rating',
            'Submit for rating'
        ]),
        ('Step 5: Upload App (Day 4-5)', [
            'Build AAB file in Android Studio',
            'Go to "Production" → "Create new release"',
            'Upload AAB file',
            'Add release notes',
            'Start rollout'
        ]),
        ('Step 6: Submit for Review (Day 5)', [
            'Review all sections',
            'Fix any issues',
            'Click "Submit for review"',
            'Wait 1-7 days',
            'App goes live'
        ])
    ]
    
    for step_title, step_items in steps:
        doc.add_heading(step_title, 2)
        for item in step_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    # 5. APPLE APP STORE - STEP BY STEP
    doc.add_heading('5. APPLE APP STORE - STEP BY STEP', 1)
    
    apple_steps = [
        ('Step 1: Create Apple ID (Day 1)', [
            'Go to https://appleid.apple.com',
            'Create new Apple ID',
            'Fill email, password',
            'Verify email'
        ]),
        ('Step 2: Join Developer Program (Day 1-2)', [
            'Go to https://developer.apple.com',
            'Sign in with Apple ID',
            'Click "Join the Apple Developer Program"',
            'Pay ₹7,500/year',
            'Fill organization details',
            'Wait for verification'
        ]),
        ('Step 3: Get D-U-N-S Number (Day 2-14)', [
            'Go to https://developer.apple.com/enroll/duns-lookup/',
            'Search your company name',
            'If not found: Request new number',
            'Wait 7-14 business days',
            'Get D-U-N-S number via email'
        ]),
        ('Step 4: Create App ID (Day 14-15)', [
            'Go to Apple Developer → Certificates, Identifiers & Profiles',
            'Click "+" to create App ID',
            'Description: "Veltrix Sports"',
            'Bundle ID: "com.veltrix.sports"',
            'Click "Register"'
        ]),
        ('Step 5: App Store Connect Setup (Day 15-16)', [
            'Go to https://appstoreconnect.apple.com',
            'Click "My Apps"',
            'Click "+" → "New App"',
            'Platform: iOS',
            'Name: "Veltrix Sports"',
            'Bundle ID: "com.veltrix.sports"'
        ]),
        ('Step 6: Complete App Information (Day 16)', [
            'App name (max 30 chars)',
            'Subtitle (max 30 chars)',
            'Description (max 4000 chars)',
            'Upload app icon (1024x1024)',
            'Upload screenshots'
        ]),
        ('Step 7: Build & Upload (Day 16-17)', [
            'Open project in Xcode',
            'Product → Archive',
            'Validate app',
            'Upload to App Store Connect'
        ]),
        ('Step 8: Submit for Review (Day 17)', [
            'Go to "App Store" tab',
            'Click "Submit for Review"',
            'Answer review questions',
            'Wait 24-48 hours'
        ])
    ]
    
    for step_title, step_items in apple_steps:
        doc.add_heading(step_title, 2)
        for item in step_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    # 6. TIMELINE COMPARISON
    doc.add_heading('6. TIMELINE COMPARISON', 1)
    
    timeline_table = doc.add_table(rows=4, cols=4)
    timeline_table.style = 'Table Grid'
    timeline_data = [
        ('Option', 'Google Play', 'Apple', 'Total Time'),
        ('Android Only', '3-10 days', '-', '3-10 days'),
        ('iOS Only', '-', '10-20 days', '10-20 days'),
        ('Both', '3-10 days', '10-20 days', '10-20 days')
    ]
    for i, row_data in enumerate(timeline_data):
        for j, cell_data in enumerate(row_data):
            timeline_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 7. COST COMPARISON
    doc.add_heading('7. COST COMPARISON', 1)
    
    cost_table = doc.add_table(rows=4, cols=4)
    cost_table.style = 'Table Grid'
    cost_data = [
        ('Option', 'Google Play', 'Apple', 'Total Cost'),
        ('Android Only', '₹18,000', '-', '₹18,000'),
        ('iOS Only', '-', '₹7,500/year', '₹7,500/year'),
        ('Both', '₹18,000', '₹7,500/year', '₹25,500 first year')
    ]
    for i, row_data in enumerate(cost_data):
        for j, cell_data in enumerate(row_data):
            cost_table.cell(i, j).text = cell_data
    
    doc.add_page_break()
    
    # 8. FINAL RECOMMENDATION
    doc.add_heading('8. FINAL RECOMMENDATION', 1)
    
    doc.add_heading('For 2-Week MVP', 2)
    mvp_table = doc.add_table(rows=4, cols=2)
    mvp_table.style = 'Table Grid'
    mvp_data = [
        ('Decision', 'Recommendation'),
        ('Store', 'Google Play Only'),
        ('Cost', '₹18,000'),
        ('Time', '3-10 days')
    ]
    for i, (key, value) in enumerate(mvp_data):
        mvp_table.cell(i, 0).text = key
        mvp_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 9. CHECKLIST
    doc.add_heading('9. CHECKLIST', 1)
    
    doc.add_heading('Android Only (Recommended)', 2)
    android_checklist = [
        'Create Google Play Developer Account (₹18,000)',
        'Create app listing',
        'Upload screenshots',
        'Add privacy policy',
        'Upload AAB file',
        'Submit for review',
        'Wait 1-7 days',
        'App goes live'
    ]
    for item in android_checklist:
        doc.add_paragraph('☐ ' + item)
    
    doc.add_heading('Both Stores', 2)
    both_checklist = [
        'Google Play (₹18,000)',
        'Apple Developer (₹7,500/year)',
        'Get D-U-N-S Number (7-14 days)',
        'Create certificates',
        'App Store Connect setup',
        'Upload via Xcode',
        'Submit for review',
        'Wait 10-20 days'
    ]
    for item in both_checklist:
        doc.add_paragraph('☐ ' + item)
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_APP_STORE_DECISION.docx')
    doc.save(output_path)
    print(f"Decision guide saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_decision_guide()
