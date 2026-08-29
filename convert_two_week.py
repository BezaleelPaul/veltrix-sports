"""
Convert 2-week plan costs to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_two_week_plan():
    doc = Document()
    
    # Title
    title = doc.add_heading('VELTRIX SPORTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('2-Week MVP Plan & Costs', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. PRODUCTION COSTS
    doc.add_heading('1. PRODUCTION COSTS (2 Weeks)', 1)
    
    doc.add_heading('1.1 One-Time Costs', 2)
    onetime_table = doc.add_table(rows=4, cols=2)
    onetime_table.style = 'Table Grid'
    onetime_data = [
        ('Item', 'Cost'),
        ('Google Play Store', '₹18,000'),
        ('Apple Developer', '₹7,500'),
        ('Domain (.com)', '₹800')
    ]
    for i, (key, value) in enumerate(onetime_data):
        onetime_table.cell(i, 0).text = key
        onetime_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('1.2 Infrastructure (2 Weeks)', 2)
    infra_table = doc.add_table(rows=5, cols=2)
    infra_table.style = 'Table Grid'
    infra_data = [
        ('Service', 'Cost'),
        ('AWS EC2', '₹0 (Free Tier)'),
        ('AWS RDS', '₹0 (Free Tier)'),
        ('AWS S3', '₹0 (Free Tier)'),
        ('Total', '₹0')
    ]
    for i, (key, value) in enumerate(infra_data):
        infra_table.cell(i, 0).text = key
        infra_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('1.3 Third-Party (2 Weeks)', 2)
    third_table = doc.add_table(rows=4, cols=2)
    third_table.style = 'Table Grid'
    third_data = [
        ('Service', 'Cost'),
        ('Razorpay', '2% per transaction'),
        ('Firebase', '₹0'),
        ('Total', '₹500 + 2%')
    ]
    for i, (key, value) in enumerate(third_data):
        third_table.cell(i, 0).text = key
        third_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 2. TOTAL COST
    doc.add_heading('2. TOTAL 2-WEEK COST', 1)
    
    total_table = doc.add_table(rows=4, cols=2)
    total_table.style = 'Table Grid'
    total_data = [
        ('Category', 'Cost'),
        ('One-Time', '₹26,300'),
        ('Infrastructure', '₹0'),
        ('Third-Party', '₹500')
    ]
    for i, (key, value) in enumerate(total_data):
        total_table.cell(i, 0).text = key
        total_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('TOTAL: ₹26,800').bold = True
    
    doc.add_page_break()
    
    # 3. WHAT YOU GET
    doc.add_heading('3. WHAT YOU GET IN 2 WEEKS', 1)
    
    doc.add_heading('Features', 2)
    features = [
        'User Registration (Email + Phone)',
        'User Login',
        'OTP Verification',
        'Dashboard',
        'Training Plans List',
        'Plan Detail View',
        'Events List',
        'Event Detail View',
        'Event Registration',
        'Basic Payment (Razorpay)'
    ]
    for feature in features:
        doc.add_paragraph('✓ ' + feature)
    
    doc.add_heading('Platforms', 2)
    platforms = [
        'Android App',
        'iOS App',
        'Web App (Optional)'
    ]
    for platform in platforms:
        doc.add_paragraph('✓ ' + platform)
    
    doc.add_page_break()
    
    # 4. TIMELINE
    doc.add_heading('4. TIMELINE', 1)
    
    timeline_table = doc.add_table(rows=11, cols=2)
    timeline_table.style = 'Table Grid'
    timeline_data = [
        ('Day', 'Focus'),
        ('Day 1', 'Backend Setup + Auth APIs'),
        ('Day 2', 'Flutter Setup + Auth Connection'),
        ('Day 3', 'Auth Screens (Login, Register, OTP)'),
        ('Day 4', 'Dashboard API + Screen'),
        ('Day 5', 'Training Plans Screen'),
        ('Day 6', 'Events Screen'),
        ('Day 7', 'Event Registration'),
        ('Day 8', 'Payment Integration'),
        ('Day 9', 'Testing + Bug Fixes'),
        ('Day 10', 'Deploy + Handover')
    ]
    for i, (key, value) in enumerate(timeline_data):
        timeline_table.cell(i, 0).text = key
        timeline_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 5. POST-LAUNCH COSTS
    doc.add_heading('5. POST-LAUNCH MONTHLY COSTS', 1)
    
    post_table = doc.add_table(rows=5, cols=2)
    post_table.style = 'Table Grid'
    post_data = [
        ('Item', 'Cost/Month'),
        ('AWS (Free Tier)', '₹0'),
        ('Domain Renewal', '₹70'),
        ('Apple Developer', '₹625 (₹7,500/12)'),
        ('Total', '~₹1,200/month')
    ]
    for i, (key, value) in enumerate(post_data):
        post_table.cell(i, 0).text = key
        post_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # 6. INVESTMENT SUMMARY
    doc.add_heading('6. INVESTMENT SUMMARY', 1)
    
    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Item', 'Amount'),
        ('2-Week Production', '₹26,800'),
        ('Total to Launch', '₹28,000')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = key
        summary_table.cell(i, 1).text = value
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Ready to launch in 2 weeks!').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'VELTRIX_SPORTS_2_WEEK_PLAN.docx')
    doc.save(output_path)
    print(f"2-week plan saved: {output_path}")
    return output_path

if __name__ == '__main__':
    create_two_week_plan()
